"""
NEXUS Omega DocProcessor — Servidor Unificado v3.1
Todo integrado: Generación + Plantillas + Conversión + Parser + OCR portable
Un solo archivo. Sin dependencias externas de módulos.
"""

import os, uuid, re, multiprocessing, sys, copy, io
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document
import pdfplumber
import uvicorn

# ── OCR opcional (se activa si Tesseract portable está presente) ──────────────
try:
    import cv2
    import numpy as np
    import pytesseract
    from PIL import Image as _PILImage

    _OCR_DISPONIBLE = True
except ImportError:
    _OCR_DISPONIBLE = False

app = FastAPI(
    title="Nexus DocProcessor v3.0",
    description="Motor unificado: generación, plantillas, conversión y parser de entrevistas.",
    version="3.0.0",
)
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"]
)

from fastapi.responses import HTMLResponse, Response

# ─────────────────────────────────────────────────────
#  DIRECTORIOS
# ─────────────────────────────────────────────────────
if getattr(sys, "frozen", False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TEMP_DIR = os.path.join(BASE_DIR, "temp_docs")
TPLS_DIR = os.path.join(BASE_DIR, "plantillas")
TESS_DIR = os.path.join(BASE_DIR, "tesseract")  # Tesseract OCR portable
TESS_EXE = os.path.join(TESS_DIR, "tesseract.exe")
for d in [TEMP_DIR, TPLS_DIR]:
    os.makedirs(d, exist_ok=True)

# Apuntar pytesseract al ejecutable portable si existe
if _OCR_DISPONIBLE and os.path.exists(TESS_EXE):
    pytesseract.pytesseract.tesseract_cmd = TESS_EXE
    os.environ["PATH"] += os.pathsep + TESS_DIR


# ─────────────────────────────────────────────────────
#  PWA — RUTAS ESTÁTICAS
# ─────────────────────────────────────────────────────
def _read_file(filename: str, mode: str = "r"):
    path = os.path.join(BASE_DIR, filename)
    if not os.path.exists(path):
        return None
    enc = "utf-8" if "b" not in mode else None
    with open(path, mode, encoding=enc) as f:
        return f.read()


@app.get("/", response_class=HTMLResponse)
async def serve_panel():
    html = _read_file("panel.html")
    if not html:
        return HTMLResponse(
            "<h1>panel.html no encontrado en la carpeta del servidor</h1>", 404
        )
    return HTMLResponse(html)


@app.get("/panel.html", response_class=HTMLResponse)
async def serve_panel_direct():
    html = _read_file("panel.html")
    return HTMLResponse(html or "<h1>No encontrado</h1>", 200 if html else 404)


@app.get("/manifest.json")
async def serve_manifest():
    content = _read_file("manifest.json")
    return Response(content or "{}", media_type="application/manifest+json")


@app.get("/sw.js")
async def serve_sw():
    content = _read_file("sw.js")
    return Response(content or "", media_type="application/javascript")


@app.get("/icon-{size}.png")
async def serve_icon(size: int):
    # Intentar PNG primero, luego SVG como fallback
    for ext in ["png", "svg"]:
        data = _read_file(f"icon-{size}.{ext}", "rb")
        if data:
            mt = "image/png" if ext == "png" else "image/svg+xml"
            return Response(data, media_type=mt)
    # Generar SVG inline como último recurso
    svg = f"""<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {size} {size}">
      <rect width="{size}" height="{size}" rx="{size//6}" fill="#070b12"/>
      <text x="{size//2}" y="{size//2}" font-family="monospace" font-size="{size//4}"
            font-weight="bold" fill="#00e5ff" text-anchor="middle" dominant-baseline="middle">NΩ</text>
    </svg>"""
    return Response(svg, media_type="image/svg+xml")


# ─────────────────────────────────────────────────────
#  PARSER DE ENTREVISTAS MINISTERIALES v4 (FGET Tabasco)
#  Maneja: narrativa libre, bullet points, formulario,
#  dictado, relato de primera persona, tercera persona
# ─────────────────────────────────────────────────────
MESES = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
}
MESES_NUM = MESES  # alias para parser de audiencias

# Palabras que indican quién hizo el daño (agresor/imputado)
_VERBOS_AGRESION = [
    "me golpeó",
    "me pegó",
    "me agredió",
    "me amenazó",
    "me jaloneó",
    "me aventó",
    "me empujó",
    "me cortó",
    "me hirió",
    "me disparó",
    "me robo",
    "me robó",
    "me extorsionó",
    "me violó",
    "me tocó",
    "me acosó",
    "me insulto",
    "me insultó",
    "me zarandeó",
    "me estranguló",
    "me quemó",
    "me pateó",
    "me mordió",
    "me arañó",
    "me rasguñó",
    "me lastimó",
    "golpeó",
    "agredió",
    "amenazó",
    "atacó",
    "disparó",
    "hirió",
    "robó",
    "asaltó",
    "acometió",
    "lesionó",
    "violó",
    "abusó",
    "lastimó",
    "intimidó",
    "acosó",
]
_DELITO_KEYWORDS = {
    "golpe": "LESIONES",
    "golpeó": "LESIONES",
    "pegó": "LESIONES",
    "pegaron": "LESIONES",
    "lesion": "LESIONES",
    "lesionó": "LESIONES",
    "lesiones": "LESIONES",
    "hirió": "LESIONES",
    "machetazo": "LESIONES CALIFICADAS",
    "cuchillo": "LESIONES CALIFICADAS",
    "navaja": "LESIONES CALIFICADAS",
    "arma": "LESIONES CALIFICADAS",
    "disparo": "LESIONES CALIFICADAS",
    "bala": "LESIONES CALIFICADAS",
    "violencia familiar": "VIOLENCIA FAMILIAR",
    "violenta": "VIOLENCIA FAMILIAR",
    "violento": "VIOLENCIA FAMILIAR",
    "patea": "VIOLENCIA FAMILIAR",
    "homicidio": "HOMICIDIO",
    "mató": "HOMICIDIO",
    "asesinó": "HOMICIDIO",
    "muerto": "HOMICIDIO",
    "feminicidio": "FEMINICIDIO",
    "robo": "ROBO",
    "robó": "ROBO",
    "asaltó": "ROBO",
    "sustrajeron": "ROBO",
    "fraude": "FRAUDE",
    "engañó": "FRAUDE",
    "estafó": "FRAUDE",
    "extorsión": "EXTORSIÓN",
    "extorsionó": "EXTORSIÓN",
    "extorsion": "EXTORSIÓN",
    "abuso sexual": "ABUSO SEXUAL",
    "tocó": "ABUSO SEXUAL",
    "acoso sexual": "ACOSO SEXUAL",
    "violación": "VIOLACIÓN",
    "violó": "VIOLACIÓN",
    "pederastia": "PEDERASTIA",
    "menor": "PEDERASTIA",
    "amenaza": "AMENAZAS",
    "amenazó": "AMENAZAS",
    "te mato": "AMENAZAS",
    "te voy a matar": "AMENAZAS",
    "secuestro": "SECUESTRO",
    "retuvo": "PRIVACIÓN ILEGAL DE LA LIBERTAD",
    "privó": "PRIVACIÓN ILEGAL DE LA LIBERTAD",
    "incumplimiento": "INCUMPLIMIENTO DE OBLIGACIONES ALIMENTARIAS",
    "alimentos": "INCUMPLIMIENTO DE OBLIGACIONES ALIMENTARIAS",
    "daño": "DAÑO EN PROPIEDAD AJENA",
    "destruyó": "DAÑO EN PROPIEDAD AJENA",
    "narcomenudeo": "NARCOMENUDEO",
    "droga": "NARCOMENUDEO",
    "portación": "PORTACIÓN DE ARMA DE FUEGO",
    "arma de fuego": "PORTACIÓN DE ARMA DE FUEGO",
}


def _limpiar_nombre(s: str) -> str:
    """Limpia un nombre extraído: quita prefijos de título, puntuación extra."""
    s = re.sub(r"\*+", "", s).strip()
    s = re.sub(
        r"^(?:C\.|SR\.|SRA\.|LIC\.|ING\.|DR\.|SEÑOR\s+|SEÑORA\s+)", "", s, flags=re.I
    ).strip()
    s = re.sub(r"\s{2,}", " ", s)
    s = s.strip(".,;:")
    return s


def _detectar_sexo(contexto: str, nombre: str = "") -> str:
    """Detecta sexo a partir del contexto narrativo."""
    cl = contexto.lower()
    masc = [
        "él ",
        "su esposo",
        "mi esposo",
        "mi marido",
        "el señor",
        "el sujeto",
        "el agresor",
        "el imputado",
        "el acusado",
        "mi hijo",
        "su hijo",
        "mi papá",
        "mi padre",
    ]
    feme = [
        "ella ",
        "su esposa",
        "mi esposa",
        "la señora",
        "la sujeta",
        "mi hija",
        "la agresora",
        "la imputada",
        "mi mamá",
        "mi madre",
        "mi pareja",
    ]
    for w in feme:
        if w in cl:
            return "Femenino"
    for w in masc:
        if w in cl:
            return "Masculino"
    return "S/D"


def parsear_entrevista(texto: str) -> Dict[str, str]:
    """
    Parser v4 — extrae campos de CUALQUIER forma de narrar una entrevista:
    - Relato libre en primera persona ("me golpeó", "mi esposo me atacó")
    - Formulario con etiquetas ("DELITO:", "IMPUTADO:", etc.)
    - Bullet points o listas
    - Dictado mezclado (parte formulario, parte narrativa)
    - Documentos Word pegados
    """
    result: Dict[str, str] = {}
    t = texto
    tl = t.lower()

    # ── PRE-PROCESO ──────────────────────────────────────────────
    t_clean = re.sub(r"\*+", " ", t)
    t_clean = re.sub(r"[ \t]{2,}", " ", t_clean)

    # ── FOLIO / NUC ──────────────────────────────────────────────
    for pat in [
        r"(?:NUC|FOLIO|CARPETA|CI[-\s]?JDM[-\s]?)[\s:\-\.]*([A-Z0-9\-\/]+)",
        r"\b(\d{3,4}/\d{4})\b",
        r"\bCI[-_]JDM[-_](\d+/\d{4})\b",
    ]:
        m = re.search(pat, t_clean, re.I)
        if m:
            val = m.group(1).strip().rstrip(".,")
            if re.search(r"\d", val):
                result["FOLIO_NUC"] = val
                break

    # ── DELITO ────────────────────────────────────────────────────
    # Primero: etiqueta explícita
    for pat in [
        r"(?:DELITO\s+PRINCIPAL|DELITO\s+FINAL|DELITO)\s*[:\-\.]+\s*([^\n,\.]{4,80})",
        r"por el delito (?:de\s+)?([A-ZÁÉÍÓÚÑ\sa-z]{4,60}?)(?:\s+cometido|\s+previsto|\s+en\s+|\.|,|\n)",
        r"se le imputa el delito de\s+([A-ZÁÉÍÓÚÑ\sa-z]{4,60}?)(?:\s+previsto|\.|,|\n)",
        r"hechos (?:que )?constituyen? el delito de\s+([A-ZÁÉÍÓÚÑ\sa-z]{4,60}?)(?:\.|,|\n)",
    ]:
        m = re.search(pat, t_clean, re.I)
        if m:
            val = m.group(1).strip().rstrip(".,").upper()
            if len(val) > 3 and not re.search(r"\b(QUE|EL|LA|UN|DEL)\b$", val):
                result["DELITO PRINCIPAL"] = val
                result["DELITO FINAL"] = val
                break

    # Si no hay etiqueta: inferir del vocabulario del relato
    if not result.get("DELITO PRINCIPAL"):
        for keyword, delito in _DELITO_KEYWORDS.items():
            if keyword in tl:
                result["DELITO PRINCIPAL"] = delito
                result["DELITO FINAL"] = delito
                break

    # ── IMPUTADO / AGRESOR ─────────────────────────────────────────
    # Lista de palabras que NO son nombres (evitar capturar el delito u otras cosas)
    _NO_NOMBRE = [
        "violencia",
        "familiar",
        "lesion",
        "homicidio",
        "delito",
        "fraude",
        "amenaza",
        "fiscal",
        "ministerio",
        "agencia",
        "juzgado",
        "tribunal",
        "carpeta",
        "causa",
        "penal",
    ]

    def _es_nombre_valido(n: str) -> bool:
        nl = n.lower().strip()
        if len(nl) < 5:
            return False
        if any(x in nl for x in _NO_NOMBRE):
            return False
        # Debe tener al menos dos palabras (nombre + apellido)
        partes = [p for p in nl.split() if len(p) > 1]
        return len(partes) >= 2

    # Primero: etiqueta explícita
    for pat in [
        r"(?:IMPUTADO|INVESTIGADO|ACUSADO|AGRESOR|SUJETO)\s*[:\-\.]+\s*([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s]{6,60}?)(?:,|\.|;|\n|ALIAS|quien)",
        # "en contra del C. NOMBRE" — requiere título para mayor precisión
        r"en contra (?:de[l]? )?(?:la\s+|el\s+)?(?:C\.|SR\.|SRA\.)\s*([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{8,50}?)(?:,|ALIAS|\.|;|\n)",
        # "en contra de NOMBRE APELLIDO" — sin título, requiere múltiples palabras
        r"en contra de\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{10,50}?)(?:,|\.|;|\n|por el delito|quien|el cual)",
        # "mi esposo/pareja Juan Carlos..."
        r"(?:mi|su)\s+(?:esposo[a]?|pareja|concubino[a]?|ex(?:-esposo[a]?)?|padre|madre|hijo[a]?|hermano[a]?|vecino[a]?|jefe|compañero[a]?)\s+(?:de\s+nombre\s+)?([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{6,50}?)(?:,|\.|;|\n|me\s+|quien)",
        r"era el C\.\s+([A-ZÁÉÍÓÚÑ\s]{8,50}?)(?:,|ALIAS|\.|;|\n)",
        r"(?:SR\.|SRA\.|C\.)\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{6,50}?)(?:,|ALIAS|\.|;|\n)",
    ]:
        m = re.search(pat, t_clean, re.I | re.MULTILINE)
        if m:
            nombre = _limpiar_nombre(m.group(1))
            victima_actual = result.get("VICTIMA ", "").upper()
            # No capturar si es igual a la víctima o no es nombre válido
            if _es_nombre_valido(nombre) and nombre.upper() not in victima_actual:
                result["IMPUTADO "] = nombre.upper()
                result["IMPUTADO"] = nombre.upper()
                break

    # Fallback informal: "[NOMBRE] me [golpeó/pegó/...]"
    if not result.get("IMPUTADO "):
        m = re.search(
            r"([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+){1,3})\s+"
            r"me\s+(?:golpeó|pegó|agredió|amenazó|atacó|hirió|cortó|disparó|lesionó|lastimó)",
            t_clean,
            re.I,
        )
        if m:
            nombre = _limpiar_nombre(m.group(1))
            if _es_nombre_valido(nombre):
                result["IMPUTADO "] = nombre.upper()
                result["IMPUTADO"] = nombre.upper()

    # Fallback relación: "mi esposo/pareja [nombre]" sin nombre a continuación de verbo
    if not result.get("IMPUTADO "):
        m = re.search(
            r"(?:mi|su)\s+(esposo[a]?|pareja|concubino[a]?|ex|padre|madre|hijo[a]?|hermano[a]?|"
            r"vecino[a]?|jefe|compañero[a]?)\s+(?:de\s+nombre\s+)?([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s]{5,50}?)"
            r"(?:\s+me\s+|\s+quien\s+|\s+el\s+cual\s+|,|\.|;|\n)",
            t_clean,
            re.I,
        )
        if m:
            relacion = m.group(1).upper()
            nombre = _limpiar_nombre(m.group(2))
            if _es_nombre_valido(nombre):
                result["IMPUTADO "] = f"{nombre.upper()} ({relacion})"
                result["IMPUTADO"] = result["IMPUTADO "]

    # Sexo e edad del imputado
    if result.get("IMPUTADO "):
        result["SEXO IMPUTADO"] = _detectar_sexo(t_clean)
        edad_m = re.search(
            r"(?:como de|aproximadamente|tiene)\s*(\d{1,2})\s*años", t_clean, re.I
        )
        result["EDAD IMPUTADO"] = edad_m.group(1) if edad_m else "S/D (MAYOR)"
        alias_m = re.search(
            r"ALIAS\s+[\"']?([A-ZÁÉÍÓÚÑA-Za-záéíóúñ\s]+?)[\"']?(?:,|\.|;|\n|quien)",
            t_clean,
            re.I,
        )
        if alias_m:
            result["ALIAS"] = alias_m.group(1).strip()

    # ── VÍCTIMA ────────────────────────────────────────────────────
    # Generalmente quien comparece = víctima. Etiqueta explícita:
    for pat in [
        r"(?:VÍCTIMA|VICTIMA|OFENDIDO[A]?|COMPARECIENTE|DENUNCIANTE|QUERELLANTE)\s*[:\-\.]+\s*([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s]{6,60}?)(?:,|\.|;|\n)",
        r"(?:de nombre|me llamo|mi nombre es|quien dijo llamarse)\s+([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s]{6,60}?)(?:,|\.|;|\n|de \d|nacido)",
        r"comparece\s+(?:ante\s+[^,]+,\s+)?(?:la\s+)?(?:C\.|SR\.|SRA\.)?\s*([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{6,50}?)(?:,|\.|;|\n|quien)",
    ]:
        m = re.search(pat, t_clean, re.I)
        if m:
            nombre = _limpiar_nombre(m.group(1))
            if len(nombre) > 5 and nombre.upper() != result.get("IMPUTADO ", ""):
                result["VICTIMA "] = nombre.upper()
                result["VICTIMA"] = nombre.upper()
                # Sexo e edad de la víctima
                result["SEXO VICTIMA"] = _detectar_sexo(t_clean[:200])
                edad_v = re.search(r"(\d{1,2})\s*años\s+de\s+edad", t_clean[:400], re.I)
                result["EDAD VICTIMA"] = edad_v.group(1) if edad_v else "S/D (MAYOR)"
                break

    if not result.get("VICTIMA "):
        result["VICTIMA "] = ""
        result["VICTIMA"] = ""
        result["SEXO VICTIMA"] = "S/D"
        result["EDAD VICTIMA"] = "S/D (MAYOR)"

    # ── FECHA ─────────────────────────────────────────────────────
    # Números escritos en texto (español legal)
    _NUMS_ES = {
        "uno": 1,
        "dos": 2,
        "tres": 3,
        "cuatro": 4,
        "cinco": 5,
        "seis": 6,
        "siete": 7,
        "ocho": 8,
        "nueve": 9,
        "diez": 10,
        "once": 11,
        "doce": 12,
        "trece": 13,
        "catorce": 14,
        "quince": 15,
        "dieciséis": 16,
        "dieciseis": 16,
        "diecisiete": 17,
        "dieciocho": 18,
        "diecinueve": 19,
        "veinte": 20,
        "veintiuno": 21,
        "veintidós": 22,
        "veintidos": 22,
        "veintitrés": 23,
        "veintitres": 23,
        "veinticuatro": 24,
        "veinticinco": 25,
        "veintiséis": 26,
        "veintiseis": 26,
        "veintisiete": 27,
        "veintiocho": 28,
        "veintinueve": 29,
        "treinta": 30,
        "treinta y uno": 31,
    }
    _ANIOS_ES = {
        "dos mil veinticuatro": 2024,
        "dos mil veinticinco": 2025,
        "dos mil veintiséis": 2026,
        "dos mil veintiseis": 2026,
        "dos mil veintisiete": 2027,
        "dos mil veintiocho": 2028,
        "dos mil veintitrés": 2023,
        "dos mil veintitres": 2023,
        "dos mil veintidós": 2022,
        "dos mil veintidos": 2022,
    }

    fecha_pats = [
        # "el día 15 de febrero del año 2026" / "el día de ayer 15 de febrero..."
        r"(?:el\s+día(?:\s+de\s+ayer)?|con\s+fecha|el\s+pasado|fecha:?)\s*(?:de\s+)?(\d{1,2})\s+de\s+(\w+)\s+del?\s+(?:año\s+)?(\d{4})",
        # "el 15 de febrero de 2026"
        r"el\s+(\d{1,2})\s+de\s+(\w+)\s+(?:de[l]?\s+)?(?:año\s+)?(\d{4})",
        # "15 de febrero de 2026" / "15 febrero 2026"
        r"\b(\d{1,2})\s+de\s+(\w+)\s+(?:de[l]?\s+)?(\d{4})\b",
        r"\b(\d{1,2})\s+(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+(\d{4})\b",
        # "FECHA: 15/02/2026" o "15/02/26"
        r"(?:FECHA|fecha)\s*[:\-\.]+\s*(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})",
        r"\b(\d{1,2})[/\-](\d{2})[/\-](\d{2,4})\b",
        # "ayer 15 feb 2026" / "ayer 15 de febrero"
        r"(?:ayer|anoche|antier)\s+(?:el\s+)?(\d{1,2})\s+(?:de\s+)?(\w+)\s+(?:de[l]?\s+)?(\d{4})",
    ]
    for pat in fecha_pats:
        m = re.search(pat, t_clean, re.I)
        if m:
            g = m.groups()
            dia = g[0]
            if re.search(r"[a-záéíóúñ]", g[1], re.I):
                mes_num = MESES.get(g[1].lower(), 0)
            else:
                mes_num = int(g[1]) if g[1].isdigit() else 0
            anio = g[2]
            if anio and len(str(anio)) == 2:
                anio = "20" + str(anio)
            if mes_num and 1 <= mes_num <= 12:
                result["FECHA DE INICIOS"] = f"{anio}-{mes_num:02d}-{int(dia):02d}"
                result["AÑO"] = str(anio)
                result["MES"] = str(mes_num)
                break

    # Fechas escritas en texto legal ("el quince de febrero de dos mil veintiséis")
    if not result.get("FECHA DE INICIOS"):
        for num_txt, num_val in _NUMS_ES.items():
            pat = rf"(?:el\s+)?{re.escape(num_txt)}\s+de\s+(\w+)\s+de\s+({'|'.join(re.escape(a) for a in _ANIOS_ES)})"
            m = re.search(pat, tl, re.I)
            if m:
                mes_num = MESES.get(m.group(1).lower(), 0)
                anio = _ANIOS_ES.get(m.group(2), 0)
                if mes_num and anio:
                    result["FECHA DE INICIOS"] = f"{anio}-{mes_num:02d}-{num_val:02d}"
                    result["AÑO"] = str(anio)
                    result["MES"] = str(mes_num)
                    break

    # ── MUNICIPIO ─────────────────────────────────────────────────
    municipios_tab = [
        "jalpa de méndez",
        "jalpa de mendez",
        "cárdenas",
        "cardenas",
        "centla",
        "comalcalco",
        "cunduacán",
        "cunduacan",
        "nacajuca",
        "paraíso",
        "paraiso",
        "huimanguillo",
        "macuspana",
        "palenque",
        "tenosique",
        "balancán",
        "balancan",
        "tacotalpa",
        "teapa",
        "jonuta",
        "emiliano zapata",
        "villahermosa",
        "centro",
        "jalpa",
        "nacajuca",
    ]
    for pat in [
        r"MUNICIPIO\s*[:\-\.]+\s*([A-ZÁÉÍÓÚÑA-Za-záéíóúñ\s]+?)(?:\s+TABASCO|\.|,|\n)",
        r"en el [Mm]unicipio de ([A-ZÁÉÍÓÚÑA-Za-záéíóúñ\s]+?)(?:\.|,|\n|\s+Tabasco)",
        r"(?:colonia|barrio|ranchería)[^,\n]+?(?:de|del)\s+([A-ZÁÉÍÓÚÑ][a-záéíóúñ\s]+?)(?:\.|,|\n)",
    ]:
        m = re.search(pat, t_clean, re.I)
        if m:
            mun = m.group(1).strip().lower()
            for mun_tab in municipios_tab:
                if mun_tab in mun:
                    result["MUNICIPIO"] = mun_tab.title()
                    break
            if result.get("MUNICIPIO"):
                break
    if not result.get("MUNICIPIO"):
        for mun_tab in municipios_tab:
            if mun_tab in tl:
                result["MUNICIPIO"] = mun_tab.title()
                break
    if not result.get("MUNICIPIO"):
        result["MUNICIPIO"] = "JALPA DE MENDEZ"

    # ── CALIFICATIVO ──────────────────────────────────────────────
    violencia_kw = [
        "machete",
        "golpe",
        "arma",
        "disparo",
        "machetazo",
        "navaja",
        "pistola",
        "golpeó",
        "agredió",
        "lesion",
        "hirió",
        "cortó",
        "pateó",
        "mordió",
        "quemó",
        "estranguló",
        "asfixió",
        "amenazó",
        "intimidó",
        "cuchillo",
    ]
    if any(v in tl for v in violencia_kw):
        result["CALIFICATIVO DEL DELITO"] = "Con violencia"
    else:
        result["CALIFICATIVO DEL DELITO"] = "Sin violencia"
    result["CONSUMACIÓN DEL DELITO"] = "Consumado"

    # ── FISCAL INICIAL ────────────────────────────────────────────
    for pat in [
        r"(?:FISCAL|AGENTE(?:\s+DEL\s+)?M\.?P\.?|REPRESENTANTE\s+SOCIAL)\s*[:\-\.]+\s*([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s\.]{6,60}?)(?:\n|,|\.|;)",
        r"(?:suscrita|suscrito|quien\s+suscribe)\s+(?:L\.?C\.?|Lic\.?|Licda\.?)\s*([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s]{6,50}?)(?:\n|,|\.|;)",
        r"(?:ante\s+(?:mí|mi)|ante\s+el\s+(?:fiscal|agente))[^,\n]+?(?:Lic[da]*\.?\s+)([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s]{6,50}?)(?:\n|,|\.|;)",
    ]:
        m = re.search(pat, t_clean, re.I)
        if m:
            val = m.group(1).strip().rstrip(".,")
            if len(val) > 5:
                result["FISCAL INICIAL"] = val
                break

    # ── TESTIGOS ──────────────────────────────────────────────────
    testigos = re.findall(
        r"(?:testigo|presenció|presencia\s+de|en\s+presencia\s+de)\s*[:\-]?\s*"
        r"(?:C\.|SR\.|SRA\.|Lic\.)?\s*([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑA-Za-záéíóúñ\s]{6,50}?)(?:,|\.|;|\n)",
        t_clean,
        re.I,
    )
    if testigos:
        result["TESTIGOS"] = " | ".join(
            _limpiar_nombre(x.strip())[:40] for x in testigos[:3] if len(x.strip()) > 5
        )

    # ── NARRATIVA DE LOS HECHOS ────────────────────────────────────
    # Tomar el primer bloque de texto largo (>60 chars) — la narración principal
    parrafos = [p.strip() for p in re.split(r"\n{2,}", t_clean) if len(p.strip()) > 60]
    if parrafos:
        # Preferir el párrafo que contenga verbos de acción
        mejor = parrafos[0]
        for p in parrafos:
            if any(v in p.lower() for v in _VERBOS_AGRESION):
                mejor = p
                break
        result["NARRATIVA DE LOS HECHOS"] = mejor[:600]
    else:
        # Si no hay párrafos largos, tomar todo el texto como narrativa
        result["NARRATIVA DE LOS HECHOS"] = t_clean[:600]

    # ── DEFAULTS ──────────────────────────────────────────────────
    defaults = {
        "VICEFISCALÍA": "Delitos Comunes",
        "FISCALÍA O AGENCIA": "JDM",
        "NACIONALIDAD VICTIMA": "Mexicana",
        "NACIONALIDAD IMPUTADO": "Mexicana",
        "¿DETENIDO?": "NO",
        "OBSERVACIONES": "",
        "INFORMACIÓN CORRECTA": "",
        "ALIAS": "",
    }
    for k, v in defaults.items():
        if k not in result:
            result[k] = v

    return result


# ─────────────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────────────
def safe_path(filename: str, directory: str) -> str:
    base = os.path.realpath(directory)
    full = os.path.realpath(os.path.join(directory, filename))
    if not full.startswith(base):
        raise HTTPException(400, "Nombre de archivo inválido")
    return full


def copy_cell_style(src, dst):
    if src.has_style:
        dst.font = copy.copy(src.font)
        dst.fill = copy.copy(src.fill)
        dst.border = copy.copy(src.border)
        dst.alignment = copy.copy(src.alignment)
        dst.number_format = src.number_format


def replace_in_paragraph(para, fields: Dict[str, str]):
    full = "".join(r.text for r in para.runs)
    new = full
    for k, v in fields.items():
        new = new.replace(k, str(v))
    if new != full and para.runs:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""


def map_inicio_to_novedad(row: dict) -> dict:
    folio = str(row.get("FOLIO_NUC", "")).strip()
    agencia = str(row.get("FISCALÍA O AGENCIA", "JDM")).strip()
    narrativa = str(row.get("NARRATIVA DE LOS HECHOS", "")).strip()
    detalle = (
        f"CI-{agencia}-{folio}.- {narrativa}" if folio and narrativa else narrativa
    )
    calif = str(row.get("CALIFICATIVO DEL DELITO", "")).strip()
    violencia = "Con violencia" if "con" in calif.lower() else "Sin violencia"
    return {
        "Desglose de delitos": str(
            row.get("DELITO FINAL", row.get("DELITO PRINCIPAL", ""))
        )
        .strip()
        .upper(),
        "Detenido": str(row.get("¿DETENIDO?", "NO")).strip().upper(),
        "Con Violencia/sin violencia": violencia,
        "Tipo de violencia": "",
        "Nombre de la victima": str(
            row.get("VICTIMA ", row.get("VICTIMA", ""))
        ).strip(),
        "Nombre del Imputado": str(
            row.get("IMPUTADO ", row.get("IMPUTADO", ""))
        ).strip(),
        "Medidas de proteccion": "",
        "Municipio": str(row.get("MUNICIPIO", "")).strip(),
        "Detalles relevantes": detalle,
    }


# ─────────────────────────────────────────────────────
#  MODELOS
# ─────────────────────────────────────────────────────
class ExcelRequest(BaseModel):
    filename: str
    sheet_name: str = "Datos"
    data: List[Dict[str, Any]]


class WordRequest(BaseModel):
    filename: str
    title: str
    paragraphs: List[str]


class FillWordRequest(BaseModel):
    template_name: str
    output_filename: str
    fields: Dict[str, str]
    table_data: Optional[List[Dict[str, Any]]] = None


class FillExcelRequest(BaseModel):
    template_name: str
    output_filename: str
    data: List[Dict[str, Any]]
    start_row: Optional[int] = None


class FillPdfRequest(BaseModel):
    template_name: str
    output_filename: str
    fields: Dict[str, str]


class IniciosToNovedadesRequest(BaseModel):
    reporte_template: str = "REPORTE_DIARIO.xlsx"
    novedades_template: str = "NOVEDADES.xlsx"
    output_filename: str = "NOVEDADES_GENERADAS"
    fecha: Optional[str] = None


class EntrevistaRequest(BaseModel):
    texto_entrevista: str
    nombre_victima: Optional[str] = None
    folio_nuc: Optional[str] = None
    fiscal: Optional[str] = None
    detenido: Optional[str] = "NO"
    reporte_template: str = "REPORTE_DIARIO.xlsx"
    output_filename: str = "REPORTE_CON_ENTREVISTA"


# ─────────────────────────────────────────────────────
#  ENDPOINTS — SALUD
# ─────────────────────────────────────────────────────
@app.get("/api/health")
async def health():
    templates = len(
        [f for f in os.listdir(TPLS_DIR) if f.endswith((".docx", ".xlsx", ".pdf"))]
    )
    return {"status": "ok", "version": "3.0.0", "templates_stored": templates}


# ─────────────────────────────────────────────────────
#  ENDPOINTS — GENERAR DESDE CERO
# ─────────────────────────────────────────────────────
@app.post("/api/generate-excel")
async def generate_excel(req: ExcelRequest):
    try:
        path = os.path.join(TEMP_DIR, f"{req.filename}_{uuid.uuid4().hex[:8]}.xlsx")
        pd.DataFrame(req.data).to_excel(path, sheet_name=req.sheet_name, index=False)
        return FileResponse(
            path,
            filename=f"{req.filename}.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/generate-word")
async def generate_word(req: WordRequest):
    try:
        path = os.path.join(TEMP_DIR, f"{req.filename}_{uuid.uuid4().hex[:8]}.docx")
        doc = Document()
        doc.add_heading(req.title, 1)
        for p in req.paragraphs:
            doc.add_paragraph(p)
        doc.save(path)
        return FileResponse(
            path,
            filename=f"{req.filename}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/extract-pdf")
async def extract_pdf(file: UploadFile = File(...)):
    """Compatibilidad: redirige al endpoint universal."""
    result = await extract_text_from_file(file)
    return {"filename": result["filename"], "raw_text": result["text"]}


@app.post("/api/extract-text-from-file")
async def extract_text_from_file(file: UploadFile = File(...)):
    """
    Extrae texto de PDF o imagen (jpg, png, bmp, tiff).
    PDF con texto nativo → pdfplumber.
    PDF escaneado o imagen → OCR con Tesseract + preprocesamiento OpenCV.
    Funciona aunque Tesseract no esté instalado (devuelve advertencia).
    """
    ext = os.path.splitext(file.filename or "")[1].lower()
    content = await file.read()
    text = ""
    warning = None

    if ext == ".pdf":
        tmp = os.path.join(TEMP_DIR, f"ocr_{uuid.uuid4().hex[:8]}.pdf")
        try:
            with open(tmp, "wb") as f:
                f.write(content)
            with pdfplumber.open(tmp) as pdf:
                paginas = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        paginas.append(t)
                text = "\n".join(paginas)

            # Si el PDF es escaneado (poco texto), intentar OCR
            if len(text.strip()) < 80:
                if _OCR_DISPONIBLE and os.path.exists(TESS_EXE):
                    warning = (
                        "PDF escaneado — aplicando OCR (puede tardar unos segundos)"
                    )
                    # Extraer imágenes de cada página con pdfplumber y aplicar OCR
                    ocr_pages = []
                    with pdfplumber.open(tmp) as pdf:
                        for page in pdf.pages:
                            img = page.to_image(resolution=200).original
                            nparr = (
                                np.frombuffer(img.tobytes(), np.uint8)
                                if hasattr(img, "tobytes")
                                else None
                            )
                            if nparr is not None:
                                cv_img = (
                                    cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                                    if nparr is not None
                                    else None
                                )
                            else:
                                # Convertir PIL a numpy
                                cv_img = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                            if cv_img is not None:
                                gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
                                thresh = cv2.adaptiveThreshold(
                                    gray,
                                    255,
                                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                    cv2.THRESH_BINARY,
                                    11,
                                    2,
                                )
                                t = pytesseract.image_to_string(thresh, lang="spa")
                                if t.strip():
                                    ocr_pages.append(t.strip())
                    text = "\n".join(ocr_pages)
                else:
                    warning = "PDF parece escaneado. Para extracción completa, sube la imagen directamente."
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"):
        if not _OCR_DISPONIBLE:
            raise HTTPException(
                400,
                "OCR no disponible. Instala NEXUS de nuevo para descargar Tesseract portable.",
            )
        if not os.path.exists(TESS_EXE):
            raise HTTPException(
                400,
                "Tesseract portable no encontrado. Ejecuta NEXUS_INICIAR.bat para instalarlo.",
            )
        try:
            nparr = np.frombuffer(content, np.uint8)
            cv_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if cv_img is None:
                pil_img = _PILImage.open(io.BytesIO(content))
                cv_img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
            # Preprocesamiento: escala de grises + umbral adaptativo
            gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
            thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            text = pytesseract.image_to_string(thresh, lang="spa")
        except Exception as e:
            raise HTTPException(400, f"Error procesando imagen: {e}")
    else:
        raise HTTPException(
            400, f"Formato '{ext}' no soportado. Usa PDF, JPG, PNG, BMP o TIFF."
        )

    return {"text": text.strip(), "warning": warning, "filename": file.filename}


# ─────────────────────────────────────────────────────
#  ENDPOINTS — PLANTILLAS
# ─────────────────────────────────────────────────────
@app.post("/api/template/upload")
async def upload_template(file: UploadFile = File(...)):
    if not any(file.filename.endswith(e) for e in (".docx", ".xlsx", ".pdf")):
        raise HTTPException(400, "Solo .docx, .xlsx, .pdf")
    try:
        dest = safe_path(file.filename, TPLS_DIR)
        content = await file.read()
        with open(dest, "wb") as f:
            f.write(content)
        analysis = await _analyze_template(dest, file.filename)
        return {"status": "ok", "filename": file.filename, "analysis": analysis}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/api/template/list")
async def list_templates():
    files = []
    for f in os.listdir(TPLS_DIR):
        if any(f.endswith(e) for e in (".docx", ".xlsx", ".pdf")):
            full = os.path.join(TPLS_DIR, f)
            files.append(
                {
                    "name": f,
                    "size_kb": round(os.path.getsize(full) / 1024, 1),
                    "type": f.rsplit(".", 1)[-1].upper(),
                }
            )
    return {"templates": files}


@app.get("/api/template/analyze/{filename}")
async def analyze_template(filename: str):
    path = safe_path(filename, TPLS_DIR)
    if not os.path.exists(path):
        raise HTTPException(404, f"'{filename}' no encontrada")
    return {"filename": filename, "analysis": await _analyze_template(path, filename)}


async def _analyze_template(path: str, filename: str) -> dict:
    analysis = {"type": "", "fields": [], "tables": [], "columns": [], "info": ""}
    try:
        if filename.endswith(".docx"):
            analysis["type"] = "word"
            doc = Document(path)
            found = set()
            for para in doc.paragraphs:
                found.update(re.findall(r"\{\{([^}]+)\}\}", para.text))
            tables_info = []
            for i, table in enumerate(doc.tables):
                headers = (
                    [c.text.strip() for c in table.rows[0].cells] if table.rows else []
                )
                tf = set()
                for row in table.rows:
                    for cell in row.cells:
                        tf.update(re.findall(r"\{\{([^}]+)\}\}", cell.text))
                tables_info.append(
                    {
                        "index": i,
                        "headers": headers,
                        "rows": len(table.rows),
                        "fields": list(tf),
                    }
                )
                found.update(tf)
            analysis["fields"] = [f"{{{{{f}}}}}" for f in sorted(found)]
            analysis["tables"] = tables_info
            analysis["info"] = (
                f"{len(doc.paragraphs)} párrafos, {len(doc.tables)} tabla(s)"
            )
        elif filename.endswith(".xlsx"):
            analysis["type"] = "excel"
            xl = pd.ExcelFile(path)
            for sheet in xl.sheet_names:
                df = xl.parse(sheet, nrows=5)
                analysis["columns"].append(
                    {
                        "sheet": sheet,
                        "headers": list(df.columns.astype(str)),
                        "sample_rows": len(df),
                    }
                )
            analysis["info"] = f"{len(xl.sheet_names)} hoja(s)"
        elif filename.endswith(".pdf"):
            analysis["type"] = "pdf"
            with pdfplumber.open(path) as pdf:
                text = "".join(p.extract_text() or "" for p in pdf.pages)
            matches = re.findall(r"\{\{([^}]+)\}\}", text)
            analysis["fields"] = [f"{{{{{m}}}}}" for m in set(matches)]
            analysis["info"] = f"PDF, {len(analysis['fields'])} campo(s) detectados"
    except Exception as e:
        analysis["info"] = f"Error: {e}"
    return analysis


@app.post("/api/template/fill-word")
async def fill_word_template(req: FillWordRequest):
    path = safe_path(req.template_name, TPLS_DIR)
    if not os.path.exists(path):
        raise HTTPException(404, f"'{req.template_name}' no encontrada")
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            replace_in_paragraph(para, req.fields)
        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_in_paragraph(para, req.fields)
            for para in section.footer.paragraphs:
                replace_in_paragraph(para, req.fields)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, req.fields)
            if req.table_data:
                headers = (
                    [c.text.strip() for c in table.rows[0].cells] if table.rows else []
                )
                if headers and set(req.table_data[0].keys()).intersection(set(headers)):
                    _fill_table(table, req.table_data)
        out = os.path.join(
            TEMP_DIR, f"{req.output_filename}_{uuid.uuid4().hex[:8]}.docx"
        )
        doc.save(out)
        return FileResponse(
            out,
            filename=f"{req.output_filename}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        raise HTTPException(500, str(e))


def _fill_table(table, data):
    if not data or len(table.rows) < 2:
        return
    template_row = table.rows[-1]
    headers = [c.text.strip() for c in table.rows[0].cells]
    for row_data in data:
        new_tr = copy.deepcopy(template_row._tr)
        table._tbl.append(new_tr)
        added = table.rows[-1]
        for i, cell in enumerate(added.cells):
            if i < len(headers):
                val = str(row_data.get(headers[i], ""))
                for para in cell.paragraphs:
                    if para.runs:
                        para.runs[0].text = val
                        [setattr(r, "text", "") for r in para.runs[1:]]


@app.post("/api/template/fill-excel")
async def fill_excel_template(req: FillExcelRequest):
    path = safe_path(req.template_name, TPLS_DIR)
    if not os.path.exists(path):
        raise HTTPException(404, f"'{req.template_name}' no encontrada")
    try:
        import openpyxl

        wb = openpyxl.load_workbook(path)
        ws = wb.active
        headers, header_row = [], None
        for row in ws.iter_rows():
            vals = [c.value for c in row if c.value is not None]
            if vals:
                header_row = row[0].row
                headers = [c.value for c in row]
                break
        if not headers:
            raise HTTPException(400, "Sin encabezados en la plantilla")
        last = header_row
        for row in ws.iter_rows(min_row=header_row + 1):
            if any(c.value is not None for c in row):
                last = row[0].row
        start = req.start_row or last + 1
        style_cells = [
            ws.cell(row=max(last, header_row + 1), column=i + 1)
            for i in range(len(headers))
        ]
        for r_idx, row_data in enumerate(req.data):
            for c_idx, header in enumerate(headers):
                if header is None:
                    continue
                val = row_data.get(str(header), row_data.get(header, ""))
                cell = ws.cell(row=start + r_idx, column=c_idx + 1, value=val)
                if c_idx < len(style_cells) and style_cells[c_idx].has_style:
                    copy_cell_style(style_cells[c_idx], cell)
        out = os.path.join(
            TEMP_DIR, f"{req.output_filename}_{uuid.uuid4().hex[:8]}.xlsx"
        )
        wb.save(out)
        return FileResponse(
            out,
            filename=f"{req.output_filename}.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


# ─────────────────────────────────────────────────────
#  ENDPOINTS — CONVERSIÓN / CICLO DOCUMENTAL
# ─────────────────────────────────────────────────────
@app.post("/api/convert/parse-entrevista")
async def parse_entrevista_endpoint(req: EntrevistaRequest):
    campos = parsear_entrevista(req.texto_entrevista)
    if req.nombre_victima:
        campos["VICTIMA "] = req.nombre_victima.upper()
        campos["VICTIMA"] = req.nombre_victima.upper()
    if req.folio_nuc:
        campos["FOLIO_NUC"] = req.folio_nuc
    if req.fiscal:
        campos["FISCAL INICIAL"] = req.fiscal
    if req.detenido:
        campos["¿DETENIDO?"] = req.detenido.upper()

    pendientes = [
        f
        for f in ["VICTIMA ", "FOLIO_NUC", "DELITO PRINCIPAL", "IMPUTADO ", "MUNICIPIO"]
        if not campos.get(f)
    ]
    return {
        "campos_detectados": campos,
        "total_detectados": len([v for v in campos.values() if v]),
        "campos_pendientes": pendientes,
        "listo_para_excel": len(pendientes) == 0,
    }


@app.post("/api/convert/entrevista-to-excel")
async def entrevista_to_excel(req: EntrevistaRequest):
    reporte_path = safe_path(req.reporte_template, TPLS_DIR)
    if not os.path.exists(reporte_path):
        raise HTTPException(
            404, f"Plantilla '{req.reporte_template}' no encontrada en /plantillas"
        )
    campos = parsear_entrevista(req.texto_entrevista)
    if req.nombre_victima:
        campos["VICTIMA "] = req.nombre_victima.upper()
        campos["VICTIMA"] = req.nombre_victima.upper()
    if req.folio_nuc:
        campos["FOLIO_NUC"] = req.folio_nuc
    if req.fiscal:
        campos["FISCAL INICIAL"] = req.fiscal
    if req.detenido:
        campos["¿DETENIDO?"] = req.detenido.upper()

    try:
        import openpyxl

        wb = openpyxl.load_workbook(reporte_path)
        if "Inicios de Carpetas" not in wb.sheetnames:
            raise HTTPException(400, "Hoja 'Inicios de Carpetas' no encontrada")
        ws = wb["Inicios de Carpetas"]
        headers = [c.value for c in list(ws.iter_rows(min_row=1, max_row=1))[0]]
        last_row = 1
        for row in ws.iter_rows(min_row=2):
            if any(c.value for c in row):
                last_row = row[0].row
        style_ref = max(last_row, 2)
        new_row = last_row + 1
        for col_idx, header in enumerate(headers, 1):
            if not header:
                continue
            val = (
                campos.get(header)
                or campos.get(str(header).strip())
                or campos.get(str(header) + " ")
                or ""
            )
            if val:
                cell = ws.cell(row=new_row, column=col_idx, value=val)
                copy_cell_style(ws.cell(row=style_ref, column=col_idx), cell)
        out = os.path.join(
            TEMP_DIR, f"{req.output_filename}_{uuid.uuid4().hex[:8]}.xlsx"
        )
        wb.save(out)
        return FileResponse(
            out,
            filename=f"{req.output_filename}.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/convert/preview-inicios")
async def preview_inicios(req: IniciosToNovedadesRequest):
    reporte_path = safe_path(req.reporte_template, TPLS_DIR)
    if not os.path.exists(reporte_path):
        raise HTTPException(404, f"'{req.reporte_template}' no encontrada")
    try:
        import openpyxl

        wb = openpyxl.load_workbook(reporte_path, data_only=True)
        if "Inicios de Carpetas" not in wb.sheetnames:
            raise HTTPException(400, "Hoja 'Inicios de Carpetas' no encontrada")
        ws = wb["Inicios de Carpetas"]
        headers = [c.value for c in list(ws.iter_rows(min_row=1, max_row=1))[0]]
        rows = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if any(v for v in row):
                d = dict(zip(headers, row))
                if req.fecha and req.fecha not in str(d.get("FECHA DE INICIOS", "")):
                    continue
                rows.append(
                    {
                        "folio": str(d.get("FOLIO_NUC", "")),
                        "delito": str(
                            d.get("DELITO FINAL", d.get("DELITO PRINCIPAL", ""))
                        ),
                        "victima": str(d.get("VICTIMA ", d.get("VICTIMA", ""))),
                        "imputado": str(d.get("IMPUTADO ", d.get("IMPUTADO", ""))),
                        "municipio": str(d.get("MUNICIPIO", "")),
                        "detenido": str(d.get("¿DETENIDO?", "NO")),
                        "narrativa": str(d.get("NARRATIVA DE LOS HECHOS", ""))[:200],
                    }
                )
        return {"total": len(rows), "rows": rows}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/convert/inicios-to-novedades")
async def inicios_to_novedades(req: IniciosToNovedadesRequest):
    reporte_path = safe_path(req.reporte_template, TPLS_DIR)
    novedades_path = safe_path(req.novedades_template, TPLS_DIR)
    for p, n in [
        (reporte_path, req.reporte_template),
        (novedades_path, req.novedades_template),
    ]:
        if not os.path.exists(p):
            raise HTTPException(404, f"'{n}' no encontrada en /plantillas")
    try:
        import openpyxl

        wb_rep = openpyxl.load_workbook(reporte_path, data_only=True)
        if "Inicios de Carpetas" not in wb_rep.sheetnames:
            raise HTTPException(400, "Hoja 'Inicios de Carpetas' no encontrada")
        ws_ini = wb_rep["Inicios de Carpetas"]
        headers = [c.value for c in list(ws_ini.iter_rows(min_row=1, max_row=1))[0]]
        rows = []
        for row in ws_ini.iter_rows(min_row=2, values_only=True):
            if any(v for v in row):
                d = dict(zip(headers, row))
                if req.fecha and req.fecha not in str(d.get("FECHA DE INICIOS", "")):
                    continue
                rows.append(d)
        if not rows:
            raise HTTPException(400, "Sin datos en Inicios de Carpetas")

        novedades_rows = [map_inicio_to_novedad(r) for r in rows]

        wb_nov = openpyxl.load_workbook(novedades_path)
        ws_nov = wb_nov["Detallado_foranea"]
        ws_nov["I1"] = len(novedades_rows)

        # Encontrar fila de encabezados de datos
        data_header_row = 3
        for i, row in enumerate(ws_nov.iter_rows(values_only=True), 1):
            if any("Desglose" in str(v) for v in row if v):
                data_header_row = i
                break

        nov_col_map = {
            "Desglose de delitos": 1,
            "Detenido": 2,
            "Con Violencia/sin violencia": 3,
            "Tipo de violencia": 4,
            "Nombre de la victima": 5,
            "Nombre del Imputado": 6,
            "Medidas de proteccion": 7,
            "Municipio": 8,
            "Detalles relevantes": 9,
        }
        style_cells = {
            c: ws_nov.cell(row=data_header_row + 1, column=c)
            for c in nov_col_map.values()
        }

        # Limpiar filas existentes
        for r in range(data_header_row + 1, ws_nov.max_row + 1):
            for c in range(1, 13):
                ws_nov.cell(row=r, column=c).value = None

        # Escribir datos
        for r_idx, nov in enumerate(novedades_rows):
            rn = data_header_row + 1 + r_idx
            for field, col_idx in nov_col_map.items():
                cell = ws_nov.cell(row=rn, column=col_idx, value=nov.get(field, ""))
                if col_idx in style_cells and style_cells[col_idx].has_style:
                    copy_cell_style(style_cells[col_idx], cell)

        out = os.path.join(
            TEMP_DIR, f"{req.output_filename}_{uuid.uuid4().hex[:8]}.xlsx"
        )
        wb_nov.save(out)
        return FileResponse(
            out,
            filename=f"{req.output_filename}.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


# ─────────────────────────────────────────────────────
#  PARSER DE REPORTES DE AUDIENCIA — v2 MULTI-FORMATO
#  Cubre 8 variantes reales identificadas en campo:
#  A) REPORTE DE AUDIENCIA .* DD DE mes YYYY (HH:MM) HORAS*
#  B) REPORTE DE AUDIENCIAS DIARIAS / FECHA: DD/MES/YYYY
#  C) FECHA Y HORA: DD DE MES DE YYYY, A LAS HH:MM HRS.
#  D) FECHA Y HORA DE AUDIENCIA: \n HH:MM.P.M. \n DEL DD DE MES
#  E) *CI-JDM-/xxx/yyyy \n Causa penal: \n DD/MM/YY
#  F) REPORTE DE AUDIENCIA DIARIA JALPA / JUICIO ORAL / DD DE MES YYYY
#  G) *REPORTE DE AUDIENCIA:* \n DD mes YYYY, HH:MM hrs.
#  H) FECHA Y HORA DE AUDIENCIA: \n HH;MM.A.M \n DD DE MES
# ─────────────────────────────────────────────────────

MESES_NUM = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
    "jan": 1,
    "feb": 2,
    "mar": 3,
    "apr": 4,
    "jun": 6,
    "jul": 7,
    "aug": 8,
    "sep": 9,
    "oct": 10,
    "nov": 11,
    "dec": 12,
}
MESES_STR = {
    "1": "ENERO",
    "2": "FEBRERO",
    "3": "MARZO",
    "4": "ABRIL",
    "5": "MAYO",
    "6": "JUNIO",
    "7": "JULIO",
    "8": "AGOSTO",
    "9": "SEPTIEMBRE",
    "10": "OCTUBRE",
    "11": "NOVIEMBRE",
    "12": "DICIEMBRE",
    "01": "ENERO",
    "02": "FEBRERO",
    "03": "MARZO",
    "04": "ABRIL",
    "05": "MAYO",
    "06": "JUNIO",
    "07": "JULIO",
    "08": "AGOSTO",
    "09": "SEPTIEMBRE",
    "10": "OCTUBRE",
    "11": "NOVIEMBRE",
    "12": "DICIEMBRE",
}

TIPOS_AUDIENCIA_MAP = {
    # exactos
    "inicial": "AUDIENCIA INICIAL",
    "audiencia inicial": "AUDIENCIA INICIAL",
    "inicial por cita": "AUDIENCIA INICIAL POR CITA",
    "inicial sin detenido": "AUDIENCIA INICIAL SIN DETENIDO",
    "intermedia": "AUDIENCIA INTERMEDIA",
    "audiencia intermedia": "AUDIENCIA INTERMEDIA",
    "intermedia (procedimiento abreviado)": "AUDIENCIA INTERMEDIA / PROC. ABREVIADO",
    "juicio": "JUICIO ORAL",
    "juicio oral": "JUICIO ORAL",
    "juicio unitario": "JUICIO ORAL",
    "cont juicio": "CONTINUACIÓN DE JUICIO",
    "cont. juicio": "CONTINUACIÓN DE JUICIO",
    "continuacion": "CONTINUACIÓN DE JUICIO",
    "continuacion de juicio": "CONTINUACIÓN DE JUICIO",
    "continuación de juicio": "CONTINUACIÓN DE JUICIO",
    "continuacion de juicio oral": "CONTINUACIÓN DE JUICIO ORAL",
    "juicio continuacion": "CONTINUACIÓN DE JUICIO",
    "juicio continuación": "CONTINUACIÓN DE JUICIO",
    "juicio cont": "CONTINUACIÓN DE JUICIO",
    "cont. juicio oral": "CONTINUACIÓN DE JUICIO ORAL",
    "abreviado": "PROCEDIMIENTO ABREVIADO",
    "procedimiento abreviado": "PROCEDIMIENTO ABREVIADO",
    "suspension": "SUSPENSIÓN CONDICIONAL DEL PROCESO",
    "suspension condicional": "SUSPENSIÓN CONDICIONAL DEL PROCESO",
    "suspension condicional del proceso": "SUSPENSIÓN CONDICIONAL DEL PROCESO",
    "suspension condicional dep proceso": "SUSPENSIÓN CONDICIONAL DEL PROCESO",
    "suspencion condicional": "SUSPENSIÓN CONDICIONAL DEL PROCESO",
    "revision": "REVISIÓN",
    "revision de suspencion condicional": "REVISIÓN DE SUSPENSIÓN CONDICIONAL",
    "revision de suspension condicional": "REVISIÓN DE SUSPENSIÓN CONDICIONAL",
    "alegatos": "ALEGATOS DE CLAUSURA",
    "alegatos de clausura": "ALEGATOS DE CLAUSURA",
    "juicio alegatos e individualizacion": "JUICIO — ALEGATOS E INDIVIDUALIZACIÓN",
    "juicio individualizacion y pago de daños": "JUICIO — INDIVIDUALIZACIÓN Y PAGO DE DAÑOS",
    "alegatos e individualizacion": "ALEGATOS E INDIVIDUALIZACIÓN",
    "individualizacion": "INDIVIDUALIZACIÓN DE SANCIONES",
    "prorroga": "PRÓRROGA DE PLAZO",
    "prorroga de plazo de cierre de investigacion": "PRÓRROGA PLAZO CIERRE DE INVESTIGACIÓN",
    "vinculacion": "AUDIENCIA INICIAL",
    "vinculacion a proceso": "VINCULACIÓN A PROCESO",
}


def _limpiar(s: str) -> str:
    """Limpia asteriscos WhatsApp/Markdown, espacios múltiples y guiones líderes."""
    s = re.sub(r"\*+", "", s)
    s = re.sub(r"^[\s\-–•]+", "", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


# Alias para compatibilidad
def _limpiar_asteriscos(s: str) -> str:
    return _limpiar(s)


def _normalizar_tipo(raw: str) -> str:
    """Normaliza el tipo de audiencia."""
    clean = _limpiar(raw).strip().rstrip(".,").lower()
    if clean in TIPOS_AUDIENCIA_MAP:
        return TIPOS_AUDIENCIA_MAP[clean]
    for k, v in TIPOS_AUDIENCIA_MAP.items():
        if k in clean or clean in k:
            return v
    return _limpiar(raw).strip().rstrip(".,").upper()


def _normalizar_carpeta(raw: str) -> str:
    """
    Normaliza el número de carpeta a formato CI-JDM-xxxx/yyyy.
    Maneja: 'CI JDM 1013/2023', 'CI-JDM- 1013/2023', 'CI-JDM-/153/2025',
            'CI-JDM_I-1323/2021', 'CPJ- VHSA-4359/2016  CI-JDM-1372/2026'
    """
    if not raw:
        return ""
    raw = _limpiar(raw).upper()
    # Si hay dos carpetas en el campo (CPJ + CI-JDM), tomar la CI-JDM
    m = re.search(r"CI[-\s_]?JDM[-\s_I]*[-/]?\s*[\d/]+", raw, re.I)
    if m:
        raw = m.group()
    # Normalizar espacios y guiones internos
    # CI JDM → CI-JDM
    raw = re.sub(r"CI\s+JDM", "CI-JDM", raw, flags=re.I)
    # CI-JDM- 1013  →  CI-JDM-1013
    raw = re.sub(r"(CI-JDM[-_I]*)\s*[-/]?\s*(\d)", r"CI-JDM-\2", raw, flags=re.I)
    # CI-JDM-/153  →  CI-JDM-153
    raw = re.sub(r"CI-JDM-/", "CI-JDM-", raw)
    # CI-JDM_I-  →  CI-JDM-
    raw = re.sub(r"CI-JDM[-_][I]?-", "CI-JDM-", raw)
    # Quitar dobles guiones
    raw = re.sub(r"-{2,}", "-", raw)
    # Tomar solo la parte CI-JDM-xxxx/yyyy
    m2 = re.search(r"CI-JDM-[\w/]+", raw)
    return m2.group().rstrip(".,") if m2 else raw.strip().rstrip(".,")


def _limpiar_fiscal(s: str) -> str:
    """Quita el municipio/institución del final del campo FISCAL."""
    s = _limpiar(s)
    # Quitar todo desde ", CENTRO DE PROCURACIÓN..." en adelante
    s = re.sub(r",?\s*CENTRO\s+DE\s+PROCUR[^\n]*", "", s, flags=re.I).strip()
    # Quitar "JALPA DE MENDEZ, TABASCO" suelto al final
    s = re.sub(r",?\s*JALPA\s+DE\s+MENDEZ[^\n]*", "", s, flags=re.I).strip()
    # Quitar asteriscos residuales y comas finales
    return s.strip().rstrip(".,*")


def _primera_linea_no_vacia(val: str) -> str:
    """Devuelve la primera línea no vacía de un bloque multilínea."""
    for ln in val.splitlines():
        ln = _limpiar(ln)
        if ln and len(ln) > 1:
            return ln
    return val.strip()


def _extraer_fecha_hora(t: str) -> tuple:
    """
    Extrae fecha y hora de cualquiera de los 8 formatos.
    Retorna (fecha_iso: str, hora_str: str)
    """
    fecha = ""
    hora = ""

    # ── HORA primero (antes de modificar t) ─────────────────
    patrones_hora = [
        r"\((\d{1,2}:\d{2})\)\s*HORAS?",  # (15:00) HORAS  — Formato A
        r"A\s+LAS\s+(\d{1,2}:\d{2})\s*HRS?",  # A LAS 09:00 HRS — Formato C
        r"(?:^|\n)\s*HORA\s*[:\*]+\s*[*\s]*(\d{1,2}:\d{2})",  # HORA: *8:30  — Formato B
        r"INICIO\s*[:\*\s]+(\d{1,2}:\d{2})\s*(?:HRS?|HORAS?)?",  # INICIO: 13:30
        r"COMENZ[OÓ]\s+A\s+LAS\s+(\d{1,2}:\d{2})",  # COMENZO A LAS 11:30 — Formato E
        r"(?:^|\n)\s*(\d{1,2}[;:]\d{2})\s*(?:A\.M|P\.M|AM|PM)",  # 08;00.A.M  — Formato H
        r"(?:^|\n)\s*(\d{1,2}:\d{2})\s*(?:A\.M|P\.M|AM|PM|HRS?|HORAS?)",  # 14:00.P.M — Formato D
        r"(\d{1,2}:\d{2})\s*(?:HRS?|HORAS?)",  # 8:00 HRS genérico
        r"(?:^|\n)(\d{1,2}:\d{2})\s*$",  # hora sola en línea
    ]
    for pat in patrones_hora:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            raw_h = m.group(1).replace(";", ":")
            hora = raw_h + (":00" if raw_h.count(":") == 1 else "")
            break

    # ── FECHA ────────────────────────────────────────────────

    # FORMATO A: "REPORTE DE AUDIENCIA .* 26 DE febrero 2026"
    m = re.search(
        r"REPORTE\s+DE\s+AUDIENCIA[^*\n]{0,40}?(\d{1,2})\s+[Dd][Ee]\s+(\w+)\s+(\d{4})",
        t,
        re.I,
    )
    if m:
        dia, mes_s, anio = m.group(1), m.group(2).lower(), m.group(3)
        mn = MESES_NUM.get(mes_s, 0)
        if mn:
            fecha = f"{anio}-{mn:02d}-{int(dia):02d}"

    # FORMATO C: "FECHA Y HORA: 26 DE FEBRERO DE 2026"
    if not fecha:
        m = re.search(
            r"FECHA\s+Y\s+HORA\s*:\s*(\d{1,2})\s+[Dd][Ee]\s+(\w+)\s+[Dd][Ee]\s+(\d{4})",
            t,
            re.I,
        )
        if m:
            dia, mes_s, anio = m.group(1), m.group(2).lower(), m.group(3)
            mn = MESES_NUM.get(mes_s, 0)
            if mn:
                fecha = f"{anio}-{mn:02d}-{int(dia):02d}"

    # FORMATO D/H: varias líneas "DEL 25 DE FEBRERO DEL 2026"
    #             o "23 de febrero del 2026" o "25 DE FEBRERO DEL 2026"
    if not fecha:
        for pat in [
            r"[Dd][Ee][Ll]?\s+(\d{1,2})\s+[Dd][Ee]\s+(\w+)\s+[Dd][Ee][Ll]?\s+(\d{4})",
            r"(?:^|\n)\s*(\d{1,2})\s+[Dd][Ee]\s+(\w+)\s+[Dd][Ee][Ll]?\s+(\d{4})",
            r"(?:día|dia)\s+(\d{1,2})\s+[Dd][Ee]\s+(\w+)\s+(?:del?\s+)?(\d{4})",
        ]:
            m = re.search(pat, t, re.I)
            if m:
                dia, mes_s, anio = m.group(1), m.group(2).lower(), m.group(3)
                mn = MESES_NUM.get(mes_s, 0)
                if mn:
                    fecha = f"{anio}-{mn:02d}-{int(dia):02d}"
                    break

    # FORMATO B: "FECHA: *25/FEBRERO/2026*"
    if not fecha:
        m = re.search(r"FECHA\s*[:\*]+\s*[*\s]*(\d{1,2})/(\w+)/(\d{4})", t, re.I)
        if m:
            dia, mes_raw, anio = m.group(1), m.group(2), m.group(3)
            mn = MESES_NUM.get(mes_raw.lower(), 0)
            if not mn and mes_raw.isdigit():
                mn = int(mes_raw)
            if mn:
                fecha = f"{anio}-{mn:02d}-{int(dia):02d}"

    # FORMATO F: "23/02/2026" o "24 DE FEBRERO 2026" (sin "DE" antes de año)
    if not fecha:
        m = re.search(r"\b(\d{1,2})/(\d{2})/(\d{4})\b", t)
        if m:
            dia, mes_n, anio = m.group(1), m.group(2), m.group(3)
            fecha = f"{anio}-{int(mes_n):02d}-{int(dia):02d}"

    # FORMATO E: "25/02/26"
    if not fecha:
        m = re.search(r"\b(\d{1,2})/(\d{2})/(\d{2})\b", t)
        if m:
            dia, mes_n, anio = m.group(1), m.group(2), "20" + m.group(3)
            fecha = f"{anio}-{int(mes_n):02d}-{int(dia):02d}"

    # FORMATO G: "23 febrero 2026, 15:30 hrs."
    if not fecha:
        m = re.search(r"(\d{1,2})\s+(\w+)\s+(\d{4})\s*,\s*(\d{1,2}:\d{2})", t, re.I)
        if m:
            dia, mes_s, anio = m.group(1), m.group(2).lower(), m.group(3)
            mn = MESES_NUM.get(mes_s, 0)
            if mn:
                fecha = f"{anio}-{mn:02d}-{int(dia):02d}"
                if not hora:
                    hora = m.group(4) + ":00"

    # FORMATO genérico: "DD DE MES YYYY" (sin DE antes del año)
    if not fecha:
        m = re.search(r"(\d{1,2})\s+[Dd][Ee]\s+(\w+)\s+(\d{4})", t, re.I)
        if m:
            dia, mes_s, anio = m.group(1), m.group(2).lower(), m.group(3)
            mn = MESES_NUM.get(mes_s, 0)
            if mn:
                fecha = f"{anio}-{mn:02d}-{int(dia):02d}"

    return fecha, hora


def parsear_reporte_audiencia(texto: str) -> Dict[str, Any]:
    """
    Parser v4 — extrae campos de los 5 formatos reales de WhatsApp/reportes JDM.

    Formatos soportados:
      A) REPORTE DE AUDIENCIA .* DD DE mes YYYY (HH:MM) ... campo: valor
      B) REPORTE DE AUDIENCIAS DIARIAS / FECHA: DD/MES/YYYY / - Campo: valor
      C) CENTRO DE PROCURACIÓN ... / FECHA Y HORA DE AUDIENCIA: / CARPETA NUMERO:
      D) CI-JDM-xxx / Causa penal:xxx / DD/MM/YY / campo: valor
      E) REPORTE DE AUDIENCIA DIARIA / JUICIO ORAL / fecha / CI-JDM...
    """
    # ── PRE-PROCESO ─────────────────────────────────────────────
    # Quitar asteriscos de WhatsApp (bold/italic), normalizar espacios
    t = re.sub(r"\*+", " ", texto)
    t = re.sub(r"[ \t]{2,}", " ", t)
    # Normalizar punto-y-coma en horas (08;00 → 08:00)
    t = re.sub(r"(\d{1,2});(\d{2})", r"\1:\2", t)
    result: Dict[str, Any] = {}

    # ── FECHA y HORA ─────────────────────────────────────────────
    fecha, hora = _extraer_fecha_hora(t)
    if fecha:
        result["FECHA "] = fecha  # con espacio para coincidir con columna Excel
    if hora:
        result["HORA"] = hora

    # ── CARPETA DE INVESTIGACIÓN ─────────────────────────────────
    carpeta_pats = [
        # "- Nº DE CARPETA: CI-JDM-263/2024"
        r"N[ÚU]?[Mm](?:ERO)?\s+DE\s+CARPETA\s*[:\-\.]+\s*([\w\s/\-_]+\d+/\d+)",
        # "CARPETA NUMERO: CI-JDM-312/2035"
        r"CARPETA\s+N[ÚU]?MERO\s*[:\-\.]+\s*([\w\s/\-_]+\d+/\d+)",
        # "CARPETA DE INVESTIGACION. - CI JDM 1013/2023"
        r"CARPETA\s+DE\s+INVES[A-Z]*[^:\n]{0,20}[:\-\.]+\s*([\w\s/\-_]+\d+/\d+)",
        # Carpeta al inicio de línea: "CI-JDM-xxx/yyyy" o "CI JDM xxx/yyyy"
        r"(?:^|\n)\s*((?:CI|CPJ)[-_\s]?(?:JDM|VHSA)[-_\s]?I?[-_\s]*/?\s*\d+/\d+)",
        # "Carpeta: CI-JDM-..."
        r"Carpeta\s*[:\-\.]+\s*([\w\s/\-_]+\d+/\d+)",
    ]
    for pat in carpeta_pats:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            raw = m.group(1).strip().split("\n")[0]
            result["CARPETA DE INV."] = _normalizar_carpeta(raw)
            break

    # ── CAUSA PENAL ──────────────────────────────────────────────
    causa_pats = [
        r"N[ÚU]MERO\s+DE\s+CAUSA\s+PENAL[^:]*[:\-\.]+\s*(\d+/\d+)",
        r"(?:CUASA|CAUSA)\s+PENAL\s*[:\-\.]+\s*(?:CP[-\s]*)?\s*(\d+/\d+)",
        r"(?:CUASA|CAUSA)\s+PENAL\s*:?\s*\n\s*(?:CP[-\s]*)?\s*(\d+/\d+)",
        r"Causa\s+penal\s*[:\-\.]\s*(\d+/\d+)",
        # "CARPETA ADMINISTRATIVA: CP 25/2025" → causa penal
        r"CARPETA\s+ADMINISTRATIVA\s*[:\-\.]+\s*CP\s+(\d+/\d+)",
    ]
    for pat in causa_pats:
        m = re.search(pat, t, re.I)
        if m:
            result["CAUSA PENAL"] = m.group(1).strip()
            break

    # ── MUNICIPIO ────────────────────────────────────────────────
    # Siempre es Jalpa de Méndez en estos reportes
    result["MUNICIPIO"] = "JALPA DE MENDEZ"

    # ── VÍCTIMA ──────────────────────────────────────────────────
    STOP_VIC = r"(?=\n\s*[-•–]?\s*(?:ASESOR|IMPUTADO|INVESTIGAD|DEFENSOR|DELITO|FISCAL|JUEZ|MNISTERIO|OBSERV|TIPO\s+DE|\Z))"
    vic_pats = [
        # con dos puntos: "VÍCTIMA U OFENDIDO: nombre"
        rf"V[IÍ]CTIMA\s+U\s+OFEND[IDA]+\s*[:\-\.]+\s*(.+?){STOP_VIC}",
        rf"V[IÍ]CTIMA\s*[:\-\.]+\s*(.+?){STOP_VIC}",
        # sin dos puntos (Formato C): "VICTIMA U OFENDIDA nombre"
        r"V[IÍ]CTIMA\s+U\s+OFEND[IDA]+\s+([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\s\.]+?)(?=\n)",
    ]
    for pat in vic_pats:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            val = _limpiar(m.group(1))
            val = _primera_linea_no_vacia(val)
            if val and val.upper() not in ("NO APLICA",):
                result["VICTIMAS (S)"] = val
            elif not val:
                continue
            if result.get("VICTIMAS (S)"):
                break
    # Fallback: "VÍCTIMA:" con valor en línea siguiente
    if not result.get("VICTIMAS (S)"):
        m = re.search(r"V[IÍ]CTIMA[^:\n]{0,20}:\s*\n\s*(.+?)(?=\n|\Z)", t, re.I)
        if m:
            val = _limpiar(m.group(1))
            if val and len(val) > 2:
                result["VICTIMAS (S)"] = val

    # ── IMPUTADO / INVESTIGADO / ENJUICIABLE ──────────────────────
    STOP_IMP = r"(?=\n\s*[-•–]?\s*(?:DEFENSOR|FISCAL|JUEZ|DELITO|OBSERV|AUDIENCIA|TIPO\s+DE|\Z))"
    imp_pats = [
        rf"(?:^|\n)\s*[-•–]?\s*IMPUTADO\s*[:\-\.]+\s*(.+?){STOP_IMP}",
        rf"INVESTIGAD[OA]\s*[:\-\.]+\s*(.+?){STOP_IMP}",
        rf"ENJUICIABLE\s*[:\-\.]+\s*(.+?){STOP_IMP}",
        r"Acusado\s*[:\-\.]+\s*(.+?)(?=\n|\Z)",
    ]
    for pat in imp_pats:
        m = re.search(pat, t, re.I | re.DOTALL | re.MULTILINE)
        if m:
            val = _limpiar(m.group(1))
            val = _primera_linea_no_vacia(val)
            if val:
                result["IMPUTADOS (S)"] = val
                break

    # ── DELITO ────────────────────────────────────────────────────
    STOP_DEL = r"(?=\n\s*[-•–]?\s*(?:INVESTIGAD|IMPUTADO|FISCAL|DEFENSOR|JUEZ|AUDIENCIA|TIPO\s+DE|OBSERV|\Z))"
    m = re.search(rf"DELITO\s*[:\-\.]+\s*(.+?){STOP_DEL}", t, re.I | re.DOTALL)
    if m:
        val = _limpiar(m.group(1))
        val = _primera_linea_no_vacia(val)
        if val:
            result["DELITO"] = val.upper()

    # ── TIPO DE AUDIENCIA ─────────────────────────────────────────
    tipo_pats = [
        # "TIPO DE AUDIENCIA: CONTINUACION DE JUICIO"  (también "AUDIENCIA:")
        r"TIPO\s+DE\s+AUDI(?:E|O)RENCIA\s*[:\-\.]*\s*\n?\s*(.+?)(?=\n|$)",
        r"(?:^|\n)\s*[-•–]?\s*TIPO\s+DE\s+AUDIENCIA\s*[:\-\.]+\s*(.+?)(?=\n|$)",
        # "AUDIENCIA: juicio individualizacion"
        r"(?:^|\n)\s*AUDIENCIA\s*[:\-\.]+\s*(.+?)(?=\n|RESOLUCI[OÓ]N|OBSERV|FISCAL|$)",
        # "Tipo de Audiencia: Alegatos"
        r"Tipo\s+de\s+Audi(?:e|o)rencia\s*[:\-\.]+\s*(.+?)(?=\n|$)",
        # standalone: "JUICIO ORAL" en línea propia
        r"(?:^|\n)\s*(JUICIO\s+ORAL|JUICIO\s+UNITARIO|CONTINUACI[OÓ]N\s+DE\s+JUICIO)\s*$",
    ]
    for pat in tipo_pats:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            raw = _limpiar(m.group(1)).rstrip(".,")
            # Descartar si contiene año (sería fecha, no tipo)
            if re.search(r"\b20\d\d\b", raw):
                continue
            # Descartar si es muy corto o inútil
            if len(raw) < 3:
                continue
            result["TIPO DE AUDIENCIA"] = _normalizar_tipo(raw)
            break
    # Fallback: "Audiencia de INICIAL"
    if "TIPO DE AUDIENCIA" not in result:
        m = re.search(r"Audiencia\s+de\s+(\w+(?:\s+\w+)?)", t, re.I)
        if m:
            result["TIPO DE AUDIENCIA"] = _normalizar_tipo(m.group(1))

    # ── RESOLUCIÓN / OBSERVACIONES ────────────────────────────────
    resol_pats = [
        # "RESOLUCION.- texto..." (Formato A)
        r"RESOLUCI[OÓ]N\.?\s*[-\.]+\s*(.+?)(?=\n\s*FISCAL|\n\s*[-•–]?\s*JUEZ|\Z)",
        # "Resolucion: texto" (Formato E)
        r"Resoluci[oó]n\s*[:\-\.]+\s*(.+?)(?=\n\s*FISCAL|\Z)",
        # "Resolucion de la Audiencia: texto"
        r"Resoluci[oó]n\s+de\s+la\s+Audiencia\s*[:\-\.]*\s*(.+?)(?=\n\s*(?:FISCAL|Es\s+cuanto)|\Z)",
        # "OBSERVACIONES: texto" — el bloque completo
        r"OBSERVACIONES\s*[:\-\.]*\s*\n?\s*(.+?)(?=\n\s*(?:REPORTE|CENTRO|FISCAL\s*:|$)|\Z)",
    ]
    for pat in resol_pats:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            val = _limpiar(m.group(1))
            # Quitar institución al final del texto
            val = re.sub(
                r",?\s*CENTRO\s+DE\s+PROCUR[^\n]*", "", val, flags=re.I
            ).strip()
            val = re.sub(r",?\s*JALPA\s+DE\s+MENDEZ[^\n]*", "", val, flags=re.I).strip()
            val = re.sub(r"\s*Es\s+cuanto\.?\s*$", "", val, flags=re.I).strip()
            val = val.rstrip(".,")
            if len(val) > 10:
                result["RESOLUCION"] = val[:1000]
                break

    # ── FISCAL ────────────────────────────────────────────────────
    fiscal_pats = [
        # "FISCAL.- LIC. Karen Ramirez López, CENTRO..." → quitar institución
        r"FISCAL\s*\.\s*-+\s*((?:LIC[DA]*\.|MD\.|DR\.|MTRO\.|LICDA\.)\s*[^\n,]+?)(?=,\s*CENTRO|,\s*\n|\n|\Z)",
        # "FISCAL QUE INTERVIENE: LIC. GERARDO..."
        r"(?:FISCAL\s+QUE\s+INTERVIENE|FISCAL\s+INTERVINIENTE)\s*[:\-\.]+\s*\n?\s*(.+?)(?=\n|\Z)",
        # "MNISTERIO PUBLICO QUE INTERVIENE: MD. ANTONIO..." (typo frecuente)
        r"MNISTERIO\s+P[ÚU]?BLIC[OA]\s+QUE\s+INTERVIENE\s*[:\-\.]+\s*\n?\s*(.+?)(?=\n|\Z)",
        # "FISCAL: LIC. Karen..."
        r"(?:^|\n)\s*[-•–]?\s*FISCAL\s*[:\-\.]+\s*((?:LIC[DA]*\.|MD\.|DR\.|MTRO\.|LICDA\.|\w{2,})\s*[^\n,]+?)(?=,\s*CENTRO|,\s*JALPA|,\s*\n|\n|\Z)",
        # genérico
        r"FISCAL\s*[:\-\.]+\s*(.{4,80}?)(?=,\s*CENTRO|,\s*JALPA|\n|\Z)",
    ]
    for pat in fiscal_pats:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            val = _limpiar_fiscal(m.group(1))
            if val and len(val) > 3:
                result["FISCAL"] = val
                break

    # ── JUEZ DE CONTROL ──────────────────────────────────────────
    juez_pats = [
        r"JUEZ\s+DE\s+CONTROL\s*[:\-\.]+\s*\n?\s*(.+?)(?=\n|\Z)",
        r"JUEZ[AZ]?\s+UNITARIO\s*[:\-\.]*\s*(.+?)(?=\n|\Z)",
        r"(?:^|\n)\s*[-•–]?\s*JUEZ[AZ]?\s*[:\-\.]+\s*(.+?)(?=\n|\Z)",
        r"(?:^|\n)\s*JUEZ[AZ]?\s+(.+?)(?=\n|\Z)",
    ]
    for pat in juez_pats:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            val = _limpiar(m.group(1)).rstrip(".,")
            if val and len(val) > 2:
                result["JUEZ"] = val
                break

    # ── DEFENSOR ─────────────────────────────────────────────────
    # Detecta PUBLICO, PARTICULAR, y también "Defensor particular NombreX"
    STOP_DEF = r"(?=\n\s*[-•–]?\s*(?:DELITO|FISCAL|JUEZ|OBSERV|TIPO\s+DE|\Z))"
    def_pats = [
        rf"DEFENSOR[AO]?\s+(?:P[ÚU]BLIC[OA]|PARTICULAR)\s*[:\-\.]+\s*\n?\s*(.+?){STOP_DEF}",
        rf"(?:^|\n)\s*DEFENSOR\s+(?:p[úu]blic[oa]\s+)?licenciad[oa]\s+(.+?)(?=\n|\Z)",
        rf"DEFENSOR\s*[:\-\.]+\s*(.+?){STOP_DEF}",
        r"Defensor\s+(?:p[úu]blic[oa]|particular)\s+(.+?)(?=\n|\Z)",
        r"Defensor\s+(.+?)(?=\n|\Z)",
    ]
    for pat in def_pats:
        m = re.search(pat, t, re.I | re.DOTALL | re.MULTILINE)
        if m:
            val = _limpiar(m.group(1))
            val = _primera_linea_no_vacia(val)
            if val and len(val) > 2:
                result["DEFENSOR"] = val
                break

    # ── ASESOR JURÍDICO ───────────────────────────────────────────
    STOP_ASE = r"(?=\n\s*[-•–]?\s*(?:IMPUTADO|INVESTIGAD|DEFENSOR|DELITO|FISCAL|\Z))"
    asesor_pats = [
        rf"ASESOR[AO]?\s+JUR[IÍ]DIC[OA]?\s*(?:P[ÚU]BLIC[OA]|PARTICULAR|GRATUITA?)?\s*[:\-\.]+\s*\n?\s*(.+?){STOP_ASE}",
        r"Asesor(?:a)?\s+jur[ií]dic[ao]?\s+(?:p[úu]blic[oa]\s+)?(?:Lic(?:enciado)?\.?\s+)?(.+?)(?=\n|\Z)",
        r"Asesor(?:a)?\s+jur[ií]dic[ao]?\s+(?:gratuita\s+)?(?:licenciad[ao]\s+)?(.+?)(?=\n|\Z)",
    ]
    for pat in asesor_pats:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            val = _limpiar(m.group(1))
            val = _primera_linea_no_vacia(val)
            if (
                val
                and len(val) > 3
                and val.upper()
                not in ("PUBLICO", "PARTICULAR", "NO APLICA", "GRATUITA")
            ):
                result["ASESOR JURÍDICO"] = val
                break

    return result


class AudienciaRequest(BaseModel):
    texto_reporte: str
    numero: Optional[int] = None  # número de fila (auto si None)
    template_name: str = "CPJ-JDM-RESULTADOS_AUDIENCIA_PLANTILLA.xlsx"
    output_filename: str = "RESULTADOS_AUDIENCIA_ACTUALIZADO"


class AudienciaParseRequest(BaseModel):
    texto_reporte: str


@app.post("/api/convert/parse-audiencia")
async def parse_audiencia_endpoint(req: AudienciaParseRequest):
    """Analiza texto de reporte de audiencia y devuelve campos estructurados."""
    campos = parsear_reporte_audiencia(req.texto_reporte)
    pendientes = [
        f
        for f in [
            "FECHA",
            "CARPETA DE INV.",
            "CAUSA PENAL",
            "VICTIMAS (S)",
            "DELITO",
            "IMPUTADOS (S)",
            "TIPO DE AUDIENCIA",
            "RESOLUCION",
            "FISCAL",
        ]
        if not campos.get(f)
    ]
    return {
        "campos_detectados": campos,
        "total_detectados": len([v for v in campos.values() if v]),
        "campos_pendientes": pendientes,
    }


@app.post("/api/convert/audiencia-to-excel")
async def audiencia_to_excel(req: AudienciaRequest):
    """
    Convierte un reporte de audiencia en texto libre a una fila
    en el Excel de Resultados de Audiencia, respetando el formato exacto.
    """
    tpl_path = safe_path(req.template_name, TPLS_DIR)
    if not os.path.exists(tpl_path):
        raise HTTPException(
            404,
            f"Plantilla '{req.template_name}' no encontrada en /plantillas.\n"
            f"Sube el archivo CPJ-JDM-RESULTADOS_AUDIENCIA_PLANTILLA.xlsx primero.",
        )

    campos = parsear_reporte_audiencia(req.texto_reporte)

    try:
        import openpyxl
        from datetime import datetime, time as dtime

        wb = openpyxl.load_workbook(tpl_path)
        ws = wb.active

        # Encontrar fila de encabezados (fila con "NÚM")
        header_row = 9  # default conocido
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            if any(str(v).strip().upper() in ("NÚM", "NUM") for v in row if v):
                header_row = i
                break

        # Encontrar última fila con datos
        last_data_row = header_row
        for row in ws.iter_rows(min_row=header_row + 1):
            if any(c.value for c in row):
                last_data_row = row[0].row

        new_row = last_data_row + 1

        # Determinar número de fila
        num = req.numero
        if num is None:
            # Contar filas de datos existentes
            count = 0
            for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
                if any(v for v in row):
                    count += 1
            num = count + 1

        # Fila de referencia para estilos
        style_ref = last_data_row if last_data_row > header_row else header_row + 1

        # Mapeo exacto de campos → columnas
        col_map = {
            "NUM": 1,  # A
            "FECHA": 2,  # B
            "HORA": 3,  # C
            "CARPETA DE INV.": 4,  # D
            "CAUSA PENAL": 5,  # E
            "MUNICIPIO": 6,  # F
            "VICTIMAS (S)": 7,  # G
            "DELITO": 8,  # H
            "IMPUTADOS (S)": 9,  # I
            "TIPO DE AUDIENCIA": 10,  # J
            "RESOLUCION": 11,  # K
            "FISCAL": 12,  # L
        }

        for field, col_idx in col_map.items():
            src_cell = ws.cell(row=style_ref, column=col_idx)
            dst_cell = ws.cell(row=new_row, column=col_idx)

            if field == "NUM":
                dst_cell.value = num
            elif field == "FECHA":
                fecha_str = campos.get("FECHA", "")
                if fecha_str:
                    try:
                        dst_cell.value = datetime.strptime(fecha_str, "%Y-%m-%d")
                        dst_cell.number_format = "DD/MM/YYYY"
                    except:
                        dst_cell.value = fecha_str
            elif field == "HORA":
                hora_str = campos.get("HORA", "")
                if hora_str:
                    try:
                        parts = hora_str.split(":")
                        dst_cell.value = dtime(int(parts[0]), int(parts[1]))
                        dst_cell.number_format = "HH:MM"
                    except:
                        dst_cell.value = hora_str
            else:
                dst_cell.value = campos.get(field, "")

            copy_cell_style(src_cell, dst_cell)

        out = os.path.join(
            TEMP_DIR, f"{req.output_filename}_{uuid.uuid4().hex[:8]}.xlsx"
        )
        wb.save(out)
        return FileResponse(
            out,
            filename=f"{req.output_filename}.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


# ─────────────────────────────────────────────────────
#  REPORTE DE AUDIENCIAS — PDF → WORD
# ─────────────────────────────────────────────────────
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parsear_pdf_audiencia(texto_pdf: str) -> Dict[str, str]:
    """
    Extrae campos de audiencia desde texto de PDF (expediente, cédula, etc.)
    Combina patrones del reporte de audiencias + datos de expediente.
    """
    t = texto_pdf
    result = {}

    # MUNICIPIO
    for pat in [
        r"CPJ[-\s]([A-ZÁÉÍÓÚÑ]+)",
        r"CENTRO DE PROCURACIÓN DE JUSTICIA DE ([A-ZÁÉÍÓÚÑ\s]+?)(?:\n|,|\.|TABASCO)",
        r"(Jalpa de [Mm]éndez|Cárdenas|Centla|Comalcalco|Palenque|Villahermosa|Macuspana|Balancán|Tenosique)",
        r"MUNICIPIO\s*[:\-]\s*([A-ZÁÉÍÓÚÑ\s]+?)(?:\n|,|\.)",
    ]:
        m = re.search(pat, t, re.I)
        if m:
            mun = m.group(1).strip()
            if "jalpa" in mun.lower():
                mun = "Jalpa de Méndez"
            result["MUNICIPIO"] = mun
            break
    if "MUNICIPIO" not in result:
        result["MUNICIPIO"] = "Jalpa de Méndez"

    # FECHA y HORA
    m = re.search(
        r"(\d{1,2})\s+de\s+(\w+)\s+(?:del?\s+)?(\d{4}).*?(\d{1,2}:\d{2})",
        t,
        re.I | re.DOTALL,
    )
    if m:
        dia, mes_str, anio, hora = (
            m.group(1),
            m.group(2).lower(),
            m.group(3),
            m.group(4),
        )
        mes_num = MESES.get(mes_str, 0)
        if mes_num:
            result["FECHA"] = f"{dia} de {m.group(2)} de {anio}"
            result["HORA"] = hora
    else:
        # Buscar fecha sola
        m = re.search(r"(\d{1,2})\s+de\s+(\w+)\s+(?:del?\s+)?(\d{4})", t, re.I)
        if m:
            result["FECHA"] = f"{m.group(1)} de {m.group(2)} de {m.group(3)}"
        m = re.search(r"(\d{1,2}:\d{2})\s*(?:horas?|hrs?)?", t, re.I)
        if m:
            result["HORA"] = m.group(1)

    # NÚM CARPETA (CI-JDM-xxx/yyyy)
    m = re.search(r"(CI[-\s]?[A-Z]{2,5}[-\s]?\d+/\d{4})", t, re.I)
    if m:
        result["NÚM DE CARPETA"] = m.group(1).strip().upper()

    # CARPETA ADMINISTRATIVA / CAUSA PENAL (xxx/yyyy sin CI-)
    patrones_causa = [
        r"CAUSA\s+PENAL\s*[:\-\.]?\s*(\d+/\d{4})",
        r"CARPETA\s+ADMINISTRATIVA\s*[:\-\.]?\s*(\d+/\d{4})",
        r"(?:CP|NUC)\s*[:\-]?\s*(\d+/\d{4})",
    ]
    for pat in patrones_causa:
        m = re.search(pat, t, re.I)
        if m:
            result["CARPETA ADMINISTRATIVA"] = m.group(1).strip()
            break
    if "CARPETA ADMINISTRATIVA" not in result:
        # Buscar segundo número de causa (distinto al de carpeta)
        nums = re.findall(r"(\d{2,4}/\d{4})", t)
        carpeta_num = re.search(r"\d+(?=/\d{4})", result.get("NÚM DE CARPETA", ""))
        for n in nums:
            if carpeta_num and n != carpeta_num.group():
                result["CARPETA ADMINISTRATIVA"] = n
                break

    # TIPO DE AUDIENCIA
    tipos = {
        "inicial": "Audiencia Inicial",
        "audiencia inicial": "Audiencia Inicial",
        "intermedia": "Audiencia Intermedia",
        "audiencia intermedia": "Audiencia Intermedia",
        "juicio oral": "Juicio Oral",
        "juicio": "Juicio Oral",
        "continuación": "Juicio en continuación",
        "cont. juicio": "Juicio en continuación",
        "abreviado": "Procedimiento Abreviado",
        "procedimiento abreviado": "Procedimiento Abreviado",
        "suspensión condicional": "Suspensión Condicional del Proceso",
        "suspension": "Suspensión Condicional del Proceso",
        "revision": "Revisión de Medidas Cautelares",
        "cateo": "Orden de Cateo",
        "vinculacion": "Audiencia Inicial",
    }
    for k, v in tipos.items():
        if k in t.lower():
            result["TIPO DE AUDIENCIA"] = v
            break

    # VÍCTIMA
    for pat in [
        r"V[IÍ]CTIMA\s*(?:U\s+OFENDIDO)?\s*[:\-\.]\s*(.+?)(?:\n|ASESOR|IMPUTADO|DEFENSOR|DELITO|FISCAL)",
        r"OFENDIDO\s*[:\-\.]\s*(.+?)(?:\n|ASESOR|IMPUTADO|DEFENSOR)",
        r"(?:DENUNCIANTE|QUERELLANTE)\s*[:\-\.]\s*(.+?)(?:\n|DELITO|IMPUTADO)",
    ]:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            result["VÍCTIMA U OFENDIDO"] = m.group(1).strip().rstrip(".,")
            break

    # ASESOR JURÍDICO
    for pat in [
        r"ASESOR\s+JUR[IÍ]DICO\s*[:\-\.]\s*(.+?)(?:\n|IMPUTADO|DEFENSOR|FISCAL)",
        r"ASESOR[AÍ]A\s*[:\-\.]\s*(.+?)(?:\n|IMPUTADO|DEFENSOR|FISCAL)",
        r"ASESOR\s*[:\-\.]\s*(.+?)(?:\n|IMPUTADO|DEFENSOR|FISCAL)",
    ]:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            result["ASESOR JURÍDICO"] = m.group(1).strip().rstrip(".,")
            break

    # IMPUTADO / INVESTIGADO
    for pat in [
        r"IMPUTADO\s*[:\-\.]\s*(.+?)(?:\n|DEFENSOR|DELITO|FISCAL|JUEZ)",
        r"INVESTIGADO\s*[:\-\.]\s*(.+?)(?:\n|DEFENSOR|DELITO|FISCAL|JUEZ)",
        r"en contra del? C\.\s+(.+?)(?:,|el cual|\n|\.)",
    ]:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            result["IMPUTADO"] = m.group(1).strip().rstrip(".,")
            break

    # DEFENSOR
    for pat in [
        r"DEFENSOR\s+(?:P[ÚU]BLICO|PARTICULAR|PRIVADO)?\s*[:\-\.]\s*(.+?)(?:\n|DELITO|FISCAL|JUEZ)",
        r"DEFENSA\s*[:\-\.]\s*(.+?)(?:\n|DELITO|FISCAL|JUEZ)",
    ]:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            result["DEFENSOR"] = m.group(1).strip().rstrip(".,")
            break

    # DELITO
    for pat in [
        r"DELITO\s*[:\-\.]\s*(.+?)(?:\n|FISCAL|JUEZ|ASESOR|DEFENSOR|OBSERV)",
        r"por el delito de\s+(.+?)(?:\n|,|\.|cometido)",
        r"DELITO\s+IMPUTADO\s*[:\-\.]\s*(.+?)(?:\n|$)",
    ]:
        m = re.search(pat, t, re.I)
        if m:
            result["DELITO"] = m.group(1).strip().rstrip(".,").upper()
            break

    # FISCAL
    for pat in [
        r"FISCAL\s*(?:QUE\s+INTERVIENE|DEL\s+MP|MINISTERIO\s+P[ÚU]BLICO)?\s*[:\-\.]\s*(.+?)(?:\n|JUEZ|$)",
        r"FISCAL\s*[:\-\.]\s*((?:LIC\.|MD\.|DR\.)\s*.+?)(?:\n|JUEZ|$)",
        r"AGENTE\s+DEL\s+(?:MINISTERIO\s+P[ÚU]BLICO|MP)\s*[:\-\.]\s*(.+?)(?:\n|JUEZ|$)",
    ]:
        m = re.search(pat, t, re.I)
        if m:
            result["FISCAL QUE INTERVIENE"] = m.group(1).strip().rstrip(".,")
            break

    # JUEZ
    for pat in [
        r"JUEZ\s*(?:DE\s+CONTROL)?\s*[:\-\.]\s*(.+?)(?:\n|OBSERV|$)",
        r"(?:ANTE\s+EL\s+)?JUEZ[AZ]?\s*[:\-\.]\s*(.+?)(?:\n|OBSERV|$)",
    ]:
        m = re.search(pat, t, re.I)
        if m:
            result["JUEZ DE CONTROL"] = m.group(1).strip().rstrip(".,")
            break

    # OBSERVACIONES / RESOLUCIÓN
    for pat in [
        r"(?:OBSERVACIONES|RESOLUCI[OÓ]N|ACUERDO)\s*[:\-\.]\s*(.+?)(?:\n\n|\Z)",
        r"(?:OBSERVACIONES|RESOLUCI[OÓ]N)\s*[:\-\.]\s*(.+?)$",
    ]:
        m = re.search(pat, t, re.I | re.DOTALL)
        if m:
            result["OBSERVACIONES"] = m.group(1).strip()[:800]
            break

    return result


def generar_word_reporte_audiencia(campos: Dict[str, str]) -> str:
    """
    Genera un documento Word con el formato oficial de Reporte de Audiencias.
    Devuelve la ruta al archivo generado.
    """
    doc = Document()

    # ── Estilos globales
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── ENCABEZADO
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("CENTRO DE PROCURACIÓN DE JUSTICIA")
    run.bold = True
    run.font.size = Pt(13)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{campos.get('MUNICIPIO','[Municipio]')}, Tabasco").bold = True

    doc.add_paragraph()  # espacio

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = titulo.add_run("REPORTE DE AUDIENCIAS DIARIAS")
    t_run.bold = True
    t_run.font.size = Pt(12)
    t_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph()

    # ── FECHA Y HORA
    fecha_str = campos.get("FECHA", "[día, mes, año]")
    hora_str = campos.get("HORA", "[hora]")
    fh = doc.add_paragraph()
    fh.add_run("FECHA Y HORA: ").bold = True
    fh.add_run(f"{fecha_str}, a las {hora_str} horas.-")

    doc.add_paragraph()

    # ── CAMPOS DEL REPORTE
    campos_formato = [
        ("Nº DE CARPETA", "NÚM DE CARPETA", "[Número de carpeta de investigación]"),
        (
            "CARPETA ADMINISTRATIVA",
            "CARPETA ADMINISTRATIVA",
            "[Número de carpeta administrativa]",
        ),
        (
            "TIPO DE AUDIENCIA",
            "TIPO DE AUDIENCIA",
            "[Inicial / Intermedia / Juicio Oral / Otra]",
        ),
        ("VÍCTIMA U OFENDIDO", "VÍCTIMA U OFENDIDO", "[Nombre completo o iniciales]"),
        ("ASESOR JURÍDICO", "ASESOR JURÍDICO", "[Nombre del asesor jurídico]"),
        ("IMPUTADO", "IMPUTADO", "[Nombre del imputado]"),
        ("DEFENSOR PARTICULAR / PÚBLICO", "DEFENSOR", "[Nombre del defensor]"),
        ("DELITO", "DELITO", "[Delito imputado]"),
        ("FISCAL QUE INTERVIENE", "FISCAL QUE INTERVIENE", "[Nombre del Fiscal]"),
        ("JUEZ DE CONTROL", "JUEZ DE CONTROL", "[Nombre del Juez de Control y región]"),
    ]

    for label, key, placeholder in campos_formato:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(4)
        run_label = p.add_run(f"- {label}: ")
        run_label.bold = True
        valor = campos.get(key, "")
        run_valor = p.add_run(valor if valor else placeholder)
        if not valor:
            run_valor.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    # ── OBSERVACIONES
    obs_title = doc.add_paragraph()
    obs_title.add_run("OBSERVACIONES:").bold = True

    obs_text = campos.get("OBSERVACIONES", "")
    obs_p = doc.add_paragraph(
        obs_text
        if obs_text
        else "[Espacio para anotar resoluciones, acuerdos, incidencias o notas relevantes]"
    )
    if not obs_text:
        obs_p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Línea de firma
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fiscal_nombre = campos.get("FISCAL QUE INTERVIENE", "_________________________")
    firma.add_run(f"{fiscal_nombre}\nFISCAL DEL MINISTERIO PÚBLICO").bold = True

    # Guardar
    out_path = os.path.join(TEMP_DIR, f"REPORTE_AUDIENCIA_{uuid.uuid4().hex[:8]}.docx")
    doc.save(out_path)
    return out_path


class PDFAudienciaRequest(BaseModel):
    output_filename: str = "REPORTE_AUDIENCIA"
    municipio: Optional[str] = None
    # Overrides manuales opcionales
    nuc: Optional[str] = None
    causa: Optional[str] = None
    tipo_audiencia: Optional[str] = None
    victima: Optional[str] = None
    asesor: Optional[str] = None
    imputado: Optional[str] = None
    defensor: Optional[str] = None
    delito: Optional[str] = None
    fiscal: Optional[str] = None
    juez: Optional[str] = None
    observaciones: Optional[str] = None
    fecha: Optional[str] = None
    hora: Optional[str] = None


@app.post("/api/audiencia/parse-pdf")
async def parse_pdf_audiencia(file: UploadFile = File(...)):
    """
    Extrae texto de un PDF de expediente/cédula y devuelve los campos
    estructurados para el reporte de audiencias.
    """
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Solo archivos PDF")
    tmp = os.path.join(TEMP_DIR, f"tmp_{uuid.uuid4().hex[:8]}.pdf")
    try:
        with open(tmp, "wb") as f:
            f.write(await file.read())
        texto = ""
        with pdfplumber.open(tmp) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texto += t + "\n"
        campos = parsear_pdf_audiencia(texto)
        pendientes = [
            k
            for k in [
                "NÚM DE CARPETA",
                "TIPO DE AUDIENCIA",
                "VÍCTIMA U OFENDIDO",
                "IMPUTADO",
                "DELITO",
                "FISCAL QUE INTERVIENE",
            ]
            if not campos.get(k)
        ]
        return {
            "texto_extraido": texto[:1000],
            "campos_detectados": campos,
            "total_detectados": len([v for v in campos.values() if v]),
            "campos_pendientes": pendientes,
        }
    except Exception as e:
        raise HTTPException(500, str(e))
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


@app.post("/api/audiencia/generar-reporte")
async def generar_reporte_audiencia(req: PDFAudienciaRequest):
    """
    Genera el Word de Reporte de Audiencias con los campos proporcionados.
    Se puede llamar con datos ya extraídos del PDF o con texto manual.
    """
    campos = {
        "MUNICIPIO": req.municipio or "Jalpa de Méndez",
        "FECHA": req.fecha or "",
        "HORA": req.hora or "",
        "NÚM DE CARPETA": req.nuc or "",
        "CARPETA ADMINISTRATIVA": req.causa or "",
        "TIPO DE AUDIENCIA": req.tipo_audiencia or "",
        "VÍCTIMA U OFENDIDO": req.victima or "",
        "ASESOR JURÍDICO": req.asesor or "",
        "IMPUTADO": req.imputado or "",
        "DEFENSOR": req.defensor or "",
        "DELITO": req.delito or "",
        "FISCAL QUE INTERVIENE": req.fiscal or "",
        "JUEZ DE CONTROL": req.juez or "",
        "OBSERVACIONES": req.observaciones or "",
    }
    try:
        out_path = generar_word_reporte_audiencia(campos)
        return FileResponse(
            out_path,
            filename=f"{req.output_filename}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/audiencia/pdf-to-reporte")
async def pdf_to_reporte_directo(
    file: UploadFile = File(...),
    output_filename: str = "REPORTE_AUDIENCIA",
    municipio: str = "Jalpa de Méndez",
):
    """
    Todo en uno: sube PDF → extrae → genera Word del reporte de audiencias.
    """
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Solo archivos PDF")
    tmp = os.path.join(TEMP_DIR, f"tmp_{uuid.uuid4().hex[:8]}.pdf")
    try:
        with open(tmp, "wb") as f:
            f.write(await file.read())
        texto = ""
        with pdfplumber.open(tmp) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texto += t + "\n"
        campos = parsear_pdf_audiencia(texto)
        if municipio:
            campos["MUNICIPIO"] = municipio
        out_path = generar_word_reporte_audiencia(campos)
        return FileResponse(
            out_path,
            filename=f"{output_filename}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        raise HTTPException(500, str(e))
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)


# ─────────────────────────────────────────────────────
#  TABLERO GENERAL — ACUMULADOR PERMANENTE
# ─────────────────────────────────────────────────────
TABLERO_PATH = os.path.join(TPLS_DIR, "TABLERO_GENERAL.xlsx")

# Estructura de cada hoja del Tablero
TABLERO_SHEETS = {
    "Inicios de Carpetas": [
        "NUC",
        "FOLIO_NUC",
        "FECHA DE INICIOS",
        "AÑO",
        "MES",
        "VICEFISCALÍA",
        "FISCALÍA O AGENCIA",
        "DELITO PRINCIPAL",
        "DELITO FINAL",
        "CALIFICATIVO DEL DELITO",
        "CONSUMACIÓN DEL DELITO",
        "¿DETENIDO?",
        "VICTIMA ",
        "SEXO VICTIMA",
        "EDAD VICTIMA",
        "NACIONALIDAD VICTIMA",
        "IMPUTADO ",
        "SEXO IMPUTADO",
        "EDAD IMPUTADO",
        "NACIONALIDAD IMPUTADO",
        "ALIAS",
        "MUNICIPIO",
        "FISCAL INICIAL",
        "NARRATIVA DE LOS HECHOS",
        "TESTIGOS",
        "OBSERVACIONES",
        "INFORMACIÓN CORRECTA",
    ],
    "Audiencias": [
        "NÚM",
        "FECHA",
        "HORA",
        "CARPETA DE INV.",
        "CAUSA PENAL",
        "MUNICIPIO",
        "VICTIMAS (S)",
        "DELITO",
        "IMPUTADOS (S)",
        "TIPO DE AUDIENCIA",
        "RESOLUCION",
        "FISCAL",
    ],
    "Novedades": [
        "Desglose de delitos",
        "Detenido",
        "Con Violencia/sin violencia",
        "Tipo de violencia",
        "Nombre de la victima",
        "Nombre del Imputado",
        "Medidas de proteccion",
        "Municipio",
        "Detalles relevantes",
    ],
}


def _crear_tablero_vacio() -> None:
    """Crea el TABLERO_GENERAL.xlsx si no existe, con todas las hojas y encabezados."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # quitar hoja default

    header_font = Font(bold=True, color="FFFFFF", size=10)
    header_fill = PatternFill("solid", fgColor="003366")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAAAAA")
    header_border = Border(left=thin, right=thin, bottom=thin)

    for sheet_name, columns in TABLERO_SHEETS.items():
        ws = wb.create_sheet(title=sheet_name)
        # Fila de encabezado en fila 1
        for col_idx, col_name in enumerate(columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = header_border
            ws.column_dimensions[cell.column_letter].width = max(14, len(col_name) + 2)
        ws.row_dimensions[1].height = 32
        ws.freeze_panes = "A2"

    # Hoja de resumen/portada
    ws_info = wb.create_sheet(title="ℹ️ Info", index=0)
    ws_info["A1"] = "TABLERO GENERAL — NEXUS Ω"
    ws_info["A1"].font = Font(bold=True, size=14, color="003366")
    ws_info["A2"] = "Acumulador permanente de todos los Reportes Diarios"
    ws_info["A4"] = "Última actualización:"
    ws_info["B4"] = "—"
    ws_info["A5"] = "Total registros Inicios:"
    ws_info["B5"] = 0
    ws_info["A6"] = "Total registros Audiencias:"
    ws_info["B6"] = 0
    ws_info.column_dimensions["A"].width = 32
    ws_info.column_dimensions["B"].width = 28

    wb.save(TABLERO_PATH)


def _get_last_data_row(ws) -> int:
    """Devuelve el número de la última fila con datos (después del encabezado)."""
    last = 1
    for row in ws.iter_rows(min_row=2):
        if any(c.value for c in row):
            last = row[0].row
    return last


def _row_key(row_dict: dict, sheet_name: str) -> str:
    """Genera una clave única por registro para detectar duplicados."""
    if sheet_name == "Inicios de Carpetas":
        return f"{row_dict.get('FOLIO_NUC','')}-{row_dict.get('DELITO PRINCIPAL','')}-{row_dict.get('VICTIMA ','')}"
    elif sheet_name == "Audiencias":
        return f"{row_dict.get('CARPETA DE INV.','')}-{row_dict.get('TIPO DE AUDIENCIA','')}-{row_dict.get('FECHA','')}"
    elif sheet_name == "Novedades":
        return f"{row_dict.get('Nombre de la victima','')}-{row_dict.get('Desglose de delitos','')}"
    return str(row_dict)


def _get_existing_keys(ws, columns: list, sheet_name: str) -> set:
    """Lee todas las filas existentes y devuelve sus claves únicas."""
    keys = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        d = dict(zip(columns, row))
        keys.add(_row_key(d, sheet_name))
    return keys


def _actualizar_tablero_hoja(
    ws_tablero, ws_fuente, sheet_name: str, columns: list, dry_run: bool = False
) -> dict:
    """
    Copia filas nuevas (no duplicadas) de ws_fuente al ws_tablero.
    Devuelve estadísticas.
    """
    import openpyxl
    from openpyxl.styles import Alignment

    # Encabezados del archivo fuente
    source_headers = []
    for cell in list(ws_fuente.iter_rows(min_row=1, max_row=1))[0]:
        source_headers.append(str(cell.value).strip() if cell.value else "")

    # Claves ya existentes en el tablero
    existing_keys = _get_existing_keys(ws_tablero, columns, sheet_name)
    last_row = _get_last_data_row(ws_tablero)

    added = 0
    skipped = 0

    for row in ws_fuente.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        # Mapear por nombre de columna
        row_dict = {}
        for col_name, val in zip(source_headers, row):
            if col_name:
                row_dict[col_name] = val

        key = _row_key(row_dict, sheet_name)
        if key in existing_keys or key == "--":
            skipped += 1
            continue

        if not dry_run:
            last_row += 1
            for col_idx, col_name in enumerate(columns, 1):
                # Buscar valor por nombre exacto o variante
                val = row_dict.get(col_name) or row_dict.get(col_name.strip()) or ""
                cell = ws_tablero.cell(row=last_row, column=col_idx, value=val)
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        existing_keys.add(key)
        added += 1

    return {"added": added, "skipped": skipped}


@app.get("/api/tablero/info")
async def tablero_info():
    """Devuelve metadatos del Tablero General."""
    import openpyxl
    from datetime import datetime

    if not os.path.exists(TABLERO_PATH):
        return {"existe": False, "sheets": {}, "ultima_actualizacion": None}

    try:
        wb = openpyxl.load_workbook(TABLERO_PATH, data_only=True)
        info = {"existe": True, "sheets": {}, "ultima_actualizacion": None}

        ws_info = wb["ℹ️ Info"] if "ℹ️ Info" in wb.sheetnames else None
        if ws_info:
            info["ultima_actualizacion"] = str(ws_info["B4"].value or "—")

        for sheet_name, columns in TABLERO_SHEETS.items():
            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                count = sum(
                    1 for row in ws.iter_rows(min_row=2, values_only=True) if any(row)
                )
                info["sheets"][sheet_name] = {
                    "total_registros": count,
                    "columnas": len(columns),
                }

        return info
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/tablero/crear")
async def crear_tablero():
    """Crea o reinicia el Tablero General vacío."""
    try:
        _crear_tablero_vacio()
        return {
            "status": "ok",
            "message": "Tablero General creado correctamente",
            "path": TABLERO_PATH,
        }
    except Exception as e:
        raise HTTPException(500, str(e))


class TableroUpdateRequest(BaseModel):
    reporte_filename: str  # nombre del archivo en /plantillas
    incluir_audiencias: bool = True
    incluir_inicios: bool = True
    incluir_novedades: bool = False  # opcional si existe hoja Novedades en el reporte


@app.post("/api/tablero/actualizar")
async def actualizar_tablero(req: TableroUpdateRequest):
    """
    Lee el Reporte Diario indicado y copia TODO su contenido
    al Tablero General, sin duplicar registros existentes.
    """
    import openpyxl
    from datetime import datetime

    reporte_path = safe_path(req.reporte_filename, TPLS_DIR)
    if not os.path.exists(reporte_path):
        raise HTTPException(
            404, f"'{req.reporte_filename}' no encontrado en /plantillas"
        )

    # Crear tablero si no existe
    if not os.path.exists(TABLERO_PATH):
        _crear_tablero_vacio()

    try:
        wb_rep = openpyxl.load_workbook(reporte_path, data_only=True)
        wb_tab = openpyxl.load_workbook(TABLERO_PATH)
        stats = {}

        # ── INICIOS DE CARPETAS
        if req.incluir_inicios and "Inicios de Carpetas" in wb_rep.sheetnames:
            ws_src = wb_rep["Inicios de Carpetas"]
            ws_dst = wb_tab["Inicios de Carpetas"]
            stats["Inicios de Carpetas"] = _actualizar_tablero_hoja(
                ws_dst,
                ws_src,
                "Inicios de Carpetas",
                TABLERO_SHEETS["Inicios de Carpetas"],
            )

        # ── AUDIENCIAS (puede llamarse de distintas formas)
        if req.incluir_audiencias:
            posibles = [
                "Audiencias",
                "Resultados de Audiencia",
                "AUDIENCIAS",
                "05 AL 09 DE MAYO DE 2025",
            ] + [s for s in wb_rep.sheetnames]
            for nombre in posibles:
                if nombre in wb_rep.sheetnames:
                    ws_src = wb_rep[nombre]
                    # Verificar que tiene columnas de audiencia
                    headers = [
                        c.value
                        for c in list(ws_src.iter_rows(min_row=1, max_row=1))[0]
                        if c.value
                    ]
                    if any("CARPETA" in str(h) or "TIPO" in str(h) for h in headers):
                        # Buscar encabezado real (puede estar en fila 9)
                        for r_idx in range(1, 12):
                            row_vals = [
                                c.value
                                for c in list(
                                    ws_src.iter_rows(min_row=r_idx, max_row=r_idx)
                                )[0]
                            ]
                            if any(
                                "NÚM" in str(v) or "CARPETA" in str(v)
                                for v in row_vals
                                if v
                            ):
                                # Crear subhoja temporal con solo los datos desde esa fila
                                import openpyxl as ox2

                                wb_tmp = ox2.Workbook()
                                ws_tmp = wb_tmp.active
                                all_rows = list(
                                    ws_src.iter_rows(min_row=r_idx, values_only=True)
                                )
                                for r_i, row in enumerate(all_rows, 1):
                                    for c_i, val in enumerate(row[:12], 1):
                                        ws_tmp.cell(row=r_i, column=c_i, value=val)
                                ws_dst = wb_tab["Audiencias"]
                                stats["Audiencias"] = _actualizar_tablero_hoja(
                                    ws_dst,
                                    ws_tmp,
                                    "Audiencias",
                                    TABLERO_SHEETS["Audiencias"],
                                )
                                break
                    break

        # ── NOVEDADES (si existe en el reporte)
        if req.incluir_novedades:
            for nombre in ["Detallado_foranea", "Novedades", "NOVEDADES"]:
                if nombre in wb_rep.sheetnames:
                    ws_src = wb_rep[nombre]
                    ws_dst = wb_tab["Novedades"]
                    stats["Novedades"] = _actualizar_tablero_hoja(
                        ws_dst, ws_src, "Novedades", TABLERO_SHEETS["Novedades"]
                    )
                    break

        # ── Actualizar Info
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        if "ℹ️ Info" in wb_tab.sheetnames:
            ws_i = wb_tab["ℹ️ Info"]
            ws_i["B4"] = now
            ws_i["B5"] = sum(
                1
                for row in wb_tab["Inicios de Carpetas"].iter_rows(
                    min_row=2, values_only=True
                )
                if any(row)
            )
            if "Audiencias" in wb_tab.sheetnames:
                ws_i["B6"] = sum(
                    1
                    for row in wb_tab["Audiencias"].iter_rows(
                        min_row=2, values_only=True
                    )
                    if any(row)
                )

        wb_tab.save(TABLERO_PATH)
        total_added = sum(v.get("added", 0) for v in stats.values())
        return {
            "status": "ok",
            "reporte": req.reporte_filename,
            "actualizado": now,
            "stats": stats,
            "total_registros_nuevos": total_added,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/api/tablero/descargar")
async def descargar_tablero():
    """Descarga el Tablero General completo."""
    if not os.path.exists(TABLERO_PATH):
        raise HTTPException(404, "Tablero General no existe aún. Créalo primero.")
    return FileResponse(
        TABLERO_PATH,
        filename="TABLERO_GENERAL_NEXUS.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/api/tablero/upload-reporte")
async def upload_y_actualizar(
    file: UploadFile = File(...),
    incluir_audiencias: bool = True,
    incluir_inicios: bool = True,
    incluir_novedades: bool = False,
):
    """
    Todo en uno: sube el Reporte Diario y actualiza el Tablero General inmediatamente.
    """
    if not file.filename.lower().endswith(".xlsx"):
        raise HTTPException(400, "Solo archivos .xlsx")
    # Guardar en plantillas
    dest = safe_path(file.filename, TPLS_DIR)
    content = await file.read()
    with open(dest, "wb") as f:
        f.write(content)
    # Actualizar tablero
    from pydantic import BaseModel as BM

    req = TableroUpdateRequest(
        reporte_filename=file.filename,
        incluir_audiencias=incluir_audiencias,
        incluir_inicios=incluir_inicios,
        incluir_novedades=incluir_novedades,
    )
    return await actualizar_tablero(req)


@app.get("/api/tablero/preview/{sheet_name}")
async def tablero_preview(sheet_name: str, limit: int = 50):
    """Devuelve las últimas filas de una hoja del Tablero para previsualización."""
    import openpyxl

    if not os.path.exists(TABLERO_PATH):
        raise HTTPException(404, "Tablero no existe")
    try:
        wb = openpyxl.load_workbook(TABLERO_PATH, data_only=True)
        if sheet_name not in wb.sheetnames:
            raise HTTPException(404, f"Hoja '{sheet_name}' no encontrada")
        ws = wb[sheet_name]
        headers = [c.value for c in list(ws.iter_rows(min_row=1, max_row=1))[0]]
        rows = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if any(row):
                rows.append(dict(zip(headers, [str(v)[:80] if v else "" for v in row])))
        # Últimas `limit` filas
        return {
            "sheet": sheet_name,
            "total": len(rows),
            "rows": rows[-limit:],
            "headers": headers,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


# ─────────────────────────────────────────────────────
#  TABLERO DE AUDIENCIAS — ACUMULADOR PERMANENTE
#  Mismo concepto que Tablero General pero para el
#  formato semanal CPJ-JDM-RESULTADOS_AUDIENCIA.xlsx
# ─────────────────────────────────────────────────────
TABLERO_AUD_PATH = os.path.join(TPLS_DIR, "TABLERO_AUDIENCIAS.xlsx")

# Columnas exactas del formato semanal de audiencias
AUD_COLUMNS = [
    "NÚM",
    "FECHA ",
    "HORA",
    "CARPETA DE INV.",
    "CAUSA PENAL",
    "MUNICIPIO",
    "VICTIMAS (S)",
    "DELITO",
    "IMPUTADOS (S)",
    "TIPO DE AUDIENCIA",
    "RESOLUCION",
    "FISCAL",
]


def _crear_tablero_audiencias() -> None:
    """
    Crea TABLERO_AUDIENCIAS.xlsx con el mismo formato visual
    que el CPJ-JDM-RESULTADOS_AUDIENCIA semanal.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "AUDIENCIAS ACUMULADAS"

    # ── ENCABEZADO (replica el formato del original)
    # Fila 1 — Título principal
    ws.merge_cells("A1:L1")
    ws["A1"] = "AUDIENCIAS DIARIAS ANTE EL JUEZ DE CONTROL"
    ws["A1"].font = Font(bold=True, size=12, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="003366")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 20

    # Fila 2 — Subtítulo
    ws.merge_cells("A2:L2")
    ws["A2"] = "DIRECCIÓN DE DELITOS COMUNES ZONA FORANEA"
    ws["A2"].font = Font(bold=True, size=11, color="003366")
    ws["A2"].fill = PatternFill("solid", fgColor="DCE6F1")
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 22

    # Fila 3 — Info tablero
    ws.merge_cells("A3:L3")
    ws["A3"] = (
        "TABLERO GENERAL DE AUDIENCIAS — NEXUS Ω  |  Registro acumulativo permanente"
    )
    ws["A3"].font = Font(italic=True, size=9, color="666666")
    ws["A3"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[3].height = 14

    # Filas 4-5 vacías
    ws.row_dimensions[4].height = 8
    ws.row_dimensions[5].height = 8

    # Fila 6 — Última actualización
    ws.merge_cells("A6:C6")
    ws["A6"] = "ÚLTIMA ACTUALIZACIÓN:"
    ws["A6"].font = Font(bold=True, size=9)
    ws.merge_cells("D6:G6")
    ws["D6"] = "—"
    ws["D6"].font = Font(size=9, color="003366")
    ws.merge_cells("H6:J6")
    ws["H6"] = "TOTAL REGISTROS:"
    ws["H6"].font = Font(bold=True, size=9)
    ws["K6"] = 0
    ws["K6"].font = Font(bold=True, size=9, color="003366")
    ws.row_dimensions[6].height = 15

    # Fila 7 — separador
    ws.merge_cells("A7:L7")
    ws["A7"].fill = PatternFill("solid", fgColor="003366")
    ws.row_dimensions[7].height = 4

    # Fila 8 vacía
    ws.row_dimensions[8].height = 6

    # Fila 9 — ENCABEZADOS DE DATOS
    header_font = Font(bold=True, size=9, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="17375E")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="AAAAAA")
    med = Side(style="medium", color="003366")

    col_widths = [6, 12, 8, 20, 14, 16, 22, 18, 22, 20, 36, 22]
    for col_idx, (col_name, width) in enumerate(zip(AUD_COLUMNS, col_widths), 1):
        cell = ws.cell(row=9, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = Border(left=thin, right=thin, top=med, bottom=med)
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.row_dimensions[9].height = 30
    ws.freeze_panes = "A10"

    wb.save(TABLERO_AUD_PATH)


def _aud_row_key(row_vals: tuple, headers: list) -> str:
    """Clave única para detectar duplicados en audiencias."""
    d = dict(zip(headers, row_vals))
    carpeta = str(d.get("CARPETA DE INV.", "")).strip()
    tipo = str(d.get("TIPO DE AUDIENCIA", "")).strip()
    fecha = str(d.get("FECHA ", "") or d.get("FECHA", "")).strip()
    fiscal = str(d.get("FISCAL", "")).strip()
    return f"{carpeta}|{tipo}|{fecha}|{fiscal}"


def _get_aud_existing_keys(ws_tab) -> set:
    """Lee claves existentes en el tablero de audiencias."""
    keys = set()
    for row in ws_tab.iter_rows(min_row=10, values_only=True):
        if not any(row):
            continue
        key = _aud_row_key(row[:12], AUD_COLUMNS)
        if key.strip("|"):
            keys.add(key)
    return keys


@app.get("/api/tablero-aud/info")
async def tablero_aud_info():
    import openpyxl

    if not os.path.exists(TABLERO_AUD_PATH):
        return {"existe": False, "total": 0, "ultima_actualizacion": "—"}
    try:
        wb = openpyxl.load_workbook(TABLERO_AUD_PATH, data_only=True)
        ws = wb.active
        total = sum(1 for row in ws.iter_rows(min_row=10, values_only=True) if any(row))
        ultima = ws["D6"].value or "—"
        return {"existe": True, "total": total, "ultima_actualizacion": str(ultima)}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/tablero-aud/crear")
async def crear_tablero_aud():
    try:
        _crear_tablero_audiencias()
        return {"status": "ok", "message": "Tablero de Audiencias creado"}
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/tablero-aud/upload")
async def upload_y_actualizar_aud(file: UploadFile = File(...)):
    """
    Sube un archivo de Resultados de Audiencia semanal (Excel)
    y copia todas las filas nuevas al Tablero de Audiencias permanente.
    """
    import openpyxl
    from openpyxl.styles import Alignment, Border, Side, Font
    from datetime import datetime as dt2

    if not file.filename.lower().endswith(".xlsx"):
        raise HTTPException(400, "Solo archivos .xlsx")

    # Guardar el archivo subido
    dest = safe_path(file.filename, TPLS_DIR)
    content = await file.read()
    with open(dest, "wb") as f:
        f.write(content)

    # Crear tablero si no existe
    if not os.path.exists(TABLERO_AUD_PATH):
        _crear_tablero_audiencias()

    try:
        wb_src = openpyxl.load_workbook(dest, data_only=True)
        wb_tab = openpyxl.load_workbook(TABLERO_AUD_PATH)
        ws_tab = wb_tab.active

        # Encontrar la hoja con datos de audiencia en el archivo fuente
        ws_src = None
        for sname in wb_src.sheetnames:
            ws_candidate = wb_src[sname]
            # Buscar fila con encabezados NÚM/CARPETA
            for r_idx in range(1, 15):
                row_vals = [
                    c.value
                    for c in list(ws_candidate.iter_rows(min_row=r_idx, max_row=r_idx))[
                        0
                    ]
                ]
                if any(str(v).strip() in ("NÚM", "NUM") for v in row_vals if v):
                    ws_src = ws_candidate
                    header_row_src = r_idx
                    break
            if ws_src:
                break

        if not ws_src:
            raise HTTPException(
                400,
                "No se encontró hoja con formato de audiencias (NÚM, CARPETA, etc.)",
            )

        # Leer encabezados fuente
        src_headers = [
            str(c.value).strip() if c.value else ""
            for c in list(
                ws_src.iter_rows(min_row=header_row_src, max_row=header_row_src)
            )[0][:12]
        ]

        # Claves ya existentes en tablero
        existing_keys = _get_aud_existing_keys(ws_tab)

        # Última fila con datos en tablero
        last_row = 9
        for row in ws_tab.iter_rows(min_row=10):
            if any(c.value for c in row):
                last_row = row[0].row

        added = 0
        skipped = 0
        thin = Side(style="thin", color="CCCCCC")
        data_border = Border(left=thin, right=thin, top=thin, bottom=thin)
        data_align = Alignment(wrap_text=True, vertical="top")

        for row in ws_src.iter_rows(min_row=header_row_src + 1, values_only=True):
            if not any(row):
                continue
            vals = list(row[:12])
            if not any(vals):
                continue

            key = _aud_row_key(tuple(vals), src_headers)
            if not key.strip("|") or key in existing_keys:
                skipped += 1
                continue

            last_row += 1
            # Mapear columna fuente → columna destino por nombre
            for dst_idx, dst_col_name in enumerate(AUD_COLUMNS, 1):
                # Buscar valor por nombre de columna
                val = ""
                for src_idx, src_col_name in enumerate(src_headers):
                    if (
                        src_col_name.strip() == dst_col_name.strip()
                        or src_col_name.strip().rstrip()
                        == dst_col_name.strip().rstrip()
                    ):
                        val = vals[src_idx] if src_idx < len(vals) else ""
                        break
                if not val and dst_idx <= len(vals):
                    val = vals[dst_idx - 1]

                cell = ws_tab.cell(row=last_row, column=dst_idx, value=val)
                cell.alignment = data_align
                cell.border = data_border
                # Alternar color de fila
                if last_row % 2 == 0:
                    from openpyxl.styles import PatternFill

                    cell.fill = PatternFill("solid", fgColor="EEF3F8")

            existing_keys.add(key)
            added += 1

        # Actualizar metadatos en encabezado
        now = dt2.now().strftime("%Y-%m-%d %H:%M")
        ws_tab["D6"] = now
        ws_tab["K6"] = sum(
            1 for row in ws_tab.iter_rows(min_row=10, values_only=True) if any(row)
        )

        wb_tab.save(TABLERO_AUD_PATH)
        return {
            "status": "ok",
            "archivo": file.filename,
            "nuevos": added,
            "ya_existian": skipped,
            "total_en_tablero": ws_tab["K6"].value,
            "actualizado": now,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/tablero-aud/agregar-texto")
async def agregar_texto_a_tablero_aud(req: AudienciaRequest):
    """
    Parsea un reporte de audiencia en texto libre y agrega
    una fila directamente al Tablero de Audiencias.
    """
    import openpyxl
    from openpyxl.styles import Alignment, Border, Side, PatternFill
    from datetime import datetime as dt2

    campos = parsear_reporte_audiencia(req.texto_reporte)

    if not os.path.exists(TABLERO_AUD_PATH):
        _crear_tablero_audiencias()

    try:
        wb = openpyxl.load_workbook(TABLERO_AUD_PATH)
        ws = wb.active

        existing_keys = _get_aud_existing_keys(ws)
        last_row = 9
        for row in ws.iter_rows(min_row=10):
            if any(c.value for c in row):
                last_row = row[0].row

        # Construir fila a partir de campos parseados
        num_actual = (
            sum(1 for row in ws.iter_rows(min_row=10, values_only=True) if any(row)) + 1
        )
        # Nota: parser v4 usa "FECHA " (con espacio) para coincidir con columna Excel
        fecha_val = campos.get("FECHA ", campos.get("FECHA", ""))
        row_data = {
            "NÚM": str(num_actual),
            "FECHA ": fecha_val,
            "HORA": campos.get("HORA", ""),
            "CARPETA DE INV.": campos.get(
                "CARPETA DE INV.", campos.get("NÚM DE CARPETA", "")
            ),
            "CAUSA PENAL": campos.get("CAUSA PENAL", ""),
            "MUNICIPIO": campos.get("MUNICIPIO", "JALPA DE MENDEZ"),
            "VICTIMAS (S)": campos.get(
                "VICTIMAS (S)", campos.get("VÍCTIMA U OFENDIDO", "")
            ),
            "DELITO": campos.get("DELITO", ""),
            "IMPUTADOS (S)": campos.get("IMPUTADOS (S)", campos.get("IMPUTADO", "")),
            "TIPO DE AUDIENCIA": campos.get("TIPO DE AUDIENCIA", ""),
            "RESOLUCION": campos.get("RESOLUCION", ""),
            "FISCAL": campos.get("FISCAL", ""),
        }

        key = f"{row_data['CARPETA DE INV.']}|{row_data['TIPO DE AUDIENCIA']}|{row_data['FECHA ']}|{row_data['FISCAL']}"
        if key in existing_keys:
            return {
                "status": "duplicado",
                "message": "Este registro ya existe en el Tablero",
            }

        thin = Side(style="thin", color="CCCCCC")
        data_border = Border(left=thin, right=thin, top=thin, bottom=thin)
        data_align = Alignment(wrap_text=True, vertical="top")

        new_row = last_row + 1
        for col_idx, col_name in enumerate(AUD_COLUMNS, 1):
            val = row_data.get(col_name, "")
            cell = ws.cell(row=new_row, column=col_idx, value=val)
            cell.alignment = data_align
            cell.border = data_border
            if new_row % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="EEF3F8")

        now = dt2.now().strftime("%Y-%m-%d %H:%M")
        ws["D6"] = now
        ws["K6"] = sum(
            1 for row in ws.iter_rows(min_row=10, values_only=True) if any(row)
        )
        wb.save(TABLERO_AUD_PATH)

        return {
            "status": "ok",
            "message": "Registro agregado al Tablero de Audiencias",
            "fila": new_row,
            "datos": row_data,
            "total": ws["K6"].value,
        }
    except Exception as e:
        raise HTTPException(500, str(e))


@app.get("/api/tablero-aud/descargar")
async def descargar_tablero_aud():
    if not os.path.exists(TABLERO_AUD_PATH):
        raise HTTPException(404, "Tablero de Audiencias no existe aún")
    return FileResponse(
        TABLERO_AUD_PATH,
        filename="TABLERO_AUDIENCIAS_NEXUS.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.get("/api/tablero-aud/preview")
async def tablero_aud_preview(limit: int = 50):
    import openpyxl

    if not os.path.exists(TABLERO_AUD_PATH):
        raise HTTPException(404, "Tablero no existe")
    try:
        wb = openpyxl.load_workbook(TABLERO_AUD_PATH, data_only=True)
        ws = wb.active
        rows = []
        for row in ws.iter_rows(min_row=10, values_only=True):
            if any(row):
                rows.append([str(v)[:60] if v else "" for v in row[:12]])
        return {"headers": AUD_COLUMNS, "rows": rows[-limit:], "total": len(rows)}
    except Exception as e:
        raise HTTPException(500, str(e))


# ─────────────────────────────────────────────────────
#  ARRANQUE
# ─────────────────────────────────────────────────────
if __name__ == "__main__":
    multiprocessing.freeze_support()
    # Railway / nube usan variable PORT; local usa 8000
    port = int(os.environ.get("PORT", 8000))
    host = "0.0.0.0" if os.environ.get("PORT") else "127.0.0.1"
    modo = "NUBE ☁️" if os.environ.get("PORT") else "LOCAL 💻"
    print("=" * 52)
    print(f"  NEXUS DOCPROCESSOR v3.0 — {modo}")
    print(f"  http://{'tu-app.railway.app' if host=='0.0.0.0' else '127.0.0.1'}:{port}")
    print("  Abre el panel en tu navegador o celular")
    print("  NO CIERRE ESTA VENTANA")
    print("=" * 52)
    uvicorn.run(app, host=host, port=port, log_level="warning")
